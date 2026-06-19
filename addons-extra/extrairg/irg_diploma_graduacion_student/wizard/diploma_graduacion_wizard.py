# -*- coding: utf-8 -*-
import os
import subprocess
import tempfile
import base64
import logging
from docx import Document as DocxDocument
from babel.dates import format_date

from odoo import models, fields, api, _
from odoo.exceptions import UserError
from odoo.modules.module import get_module_resource

_logger = logging.getLogger(__name__)


class IrgDiplomaGraduacionWizard(models.TransientModel):
    _name = 'irg.diploma.graduacion.wizard'
    _description = 'Asistente para Imprimir Diploma de Graduación'

    student_id = fields.Many2one(
        'op.student',
        string='Estudiante',
        required=True,
        readonly=True
    )
    student_course_id = fields.Many2one(
        'op.student.course',
        string='Curso Académico',
        required=True,
        domain="[('student_id', '=', student_id)]"
    )
    date = fields.Date(
        string='Fecha de Expedición',
        default=fields.Date.context_today,
        required=True
    )

    @api.model
    def default_get(self, fields_list):
        res = super(IrgDiplomaGraduacionWizard, self).default_get(fields_list)
        if 'student_id' in fields_list and not res.get('student_id'):
            res['student_id'] = self.env.context.get('active_id') or self.env.context.get('default_student_id')
        return res

    def _normalize_catalan_course_name(self, course_name):
        """Normalize common accent differences for Catalan rendering."""
        if not course_name:
            return ""
        normalized = course_name
        normalized = normalized.replace("Máster", "Màster")
        normalized = normalized.replace("máster", "màster")
        normalized = normalized.replace("Master", "Màster")
        normalized = normalized.replace("master", "màster")
        normalized = normalized.replace("Salud", "Salut")
        normalized = normalized.replace("salud", "salut")
        normalized = normalized.replace(" y ", " i ")
        normalized = normalized.replace(" Y ", " I ")
        return normalized

    @staticmethod
    def _replace_in_paragraph(paragraph, replacements):
        """Replace placeholders in a paragraph across runs that may be split by Word."""
        if not paragraph.runs:
            return
        full = ''.join(r.text for r in paragraph.runs)
        replaced = False
        for old, new in replacements.items():
            if old in full:
                full = full.replace(old, new or '')
                replaced = True
        if replaced:
            paragraph.runs[0].text = full
            for r in paragraph.runs[1:]:
                for t in r._element.xpath('.//w:t'):
                    t.text = ''

    def action_print_pdf(self):
        self.ensure_one()

        # Format dates using Babel
        try:
            date_es = "{} de {} de {}".format(
                self.date.day,
                format_date(self.date, format='MMMM', locale='es_ES'),
                self.date.year
            )
            date_cat = "{} de {} de {}".format(
                self.date.day,
                format_date(self.date, format='MMMM', locale='ca_ES'),
                self.date.year
            )
        except Exception as e:
            _logger.warning("Babel date format failed, falling back to simple strftime: %s", e)
            months_es = {
                1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril', 5: 'mayo', 6: 'junio',
                7: 'julio', 8: 'agosto', 9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
            }
            months_cat = {
                1: 'gener', 2: 'febrer', 3: 'març', 4: 'abril', 5: 'maig', 6: 'juny',
                7: 'juliol', 8: 'agost', 9: 'setembre', 10: 'octubre', 11: 'novembre', 12: 'desembre'
            }
            date_es = f"{self.date.day} de {months_es.get(self.date.month, '')} de {self.date.year}"
            date_cat = f"{self.date.day} de {months_cat.get(self.date.month, '')} de {self.date.year}"

        # Get values
        student_name = self.student_id.name or ""
        course = self.student_course_id.course_id
        course_name_es = course.name or ""
        
        course_name_cat = course.name_cat if 'name_cat' in course._fields else course_name_es
        if not course_name_cat:
            course_name_cat = course_name_es
        course_name_cat = self._normalize_catalan_course_name(course_name_cat)

        replacements = {
            '<<Alumno>>': student_name,
            '<<alumno>>': student_name,
            '<<Master>>': course_name_es,
            '<<curso>>': course_name_es,
            '<<MasterCat>>': course_name_cat,
            '<<Fecha>>': date_es,
            '<<FechaCat>>': date_cat,
        }

        tpl_path = get_module_resource(
            'irg_diploma_graduacion_student',
            'static',
            'src',
            'templates',
            'plantilla_diploma_graduado.docx'
        )

        if not tpl_path or not os.path.exists(tpl_path):
            raise UserError(_("No se encontró la plantilla Word en la ruta: %s") % tpl_path)

        # Process document using python-docx
        try:
            doc = DocxDocument(tpl_path)
        except Exception as e:
            raise UserError(_("Error al abrir la plantilla Word: %s") % str(e))

        # 1. Replace in normal paragraphs
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replacements)

        # 2. Replace in tables
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)

        # 3. Replace in headers/footers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    self._replace_in_paragraph(para, replacements)
                for tbl in section.header.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                self._replace_in_paragraph(para, replacements)
            if section.footer:
                for para in section.footer.paragraphs:
                    self._replace_in_paragraph(para, replacements)
                for tbl in section.footer.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                self._replace_in_paragraph(para, replacements)

        # Save to temp file
        tmp_docx = tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False, prefix='diploma_graduado_'
        )
        doc.save(tmp_docx.name)
        tmp_docx.close()

        # Convert to PDF
        out_dir = os.path.dirname(tmp_docx.name)
        try:
            subprocess.run(
                [
                    'libreoffice', '--headless', '--norestore',
                    '--convert-to', 'pdf',
                    '--outdir', out_dir,
                    tmp_docx.name,
                ],
                check=True,
                timeout=60,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
        except FileNotFoundError:
            if os.path.exists(tmp_docx.name):
                os.unlink(tmp_docx.name)
            raise UserError(
                _('LibreOffice no está instalado en el servidor. '
                  'Ejecute: apt-get install -y libreoffice-writer')
            )
        except subprocess.CalledProcessError as exc:
            if os.path.exists(tmp_docx.name):
                os.unlink(tmp_docx.name)
            _logger.error('LibreOffice conversion failed: %s', exc.stderr)
            raise UserError(
                _('Error al convertir el diploma a PDF. Revise el log de error.')
            )

        pdf_path = tmp_docx.name.rsplit('.', 1)[0] + '.pdf'
        if not os.path.isfile(pdf_path):
            if os.path.exists(tmp_docx.name):
                os.unlink(tmp_docx.name)
            raise UserError(_('No se generó el archivo PDF.'))

        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()

        # Cleanup temp files
        for path in (tmp_docx.name, pdf_path):
            try:
                if os.path.exists(path):
                    os.unlink(path)
            except Exception:
                pass

        # Create attachment in Odoo
        filename = "Diploma_{}.pdf".format(
            student_name.replace(' ', '_')
        )
        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(pdf_bytes),
            'res_model': 'op.student',
            'res_id': self.student_id.id,
            'mimetype': 'application/pdf',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': '/web/content/%s?download=true' % attachment.id,
            'target': 'new',
        }
