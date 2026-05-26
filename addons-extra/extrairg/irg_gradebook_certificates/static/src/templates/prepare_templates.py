import os
import shutil
import sys
import docx

def main():
    if len(sys.argv) > 2:
        src_dir = sys.argv[1]
        dst_dir = sys.argv[2]
    else:
        src_dir = os.path.dirname(os.path.abspath(__file__))
        dst_dir = src_dir

    print(f"Source dir: {src_dir}")
    print(f"Destination dir: {dst_dir}")

    # Define source paths
    raimon_src = os.path.join(src_dir, "Plantilla-certificado-notas-raimon.docx")
    dpto_src = os.path.join(src_dir, "Plantilla-certificado-notas-dpto.docx")
    
    # Destination paths
    attendance_raimon = os.path.join(dst_dir, "Plantilla-certificado-asistencia-raimon.docx")
    enrollment_raimon = os.path.join(dst_dir, "Plantilla-certificado-curso-raimon.docx")
    
    attendance_dpto = os.path.join(dst_dir, "Plantilla-certificado-asistencia-dpto.docx")
    enrollment_dpto = os.path.join(dst_dir, "Plantilla-certificado-curso-dpto.docx")
    
    # Copy
    shutil.copy2(raimon_src, attendance_raimon)
    shutil.copy2(raimon_src, enrollment_raimon)
    shutil.copy2(dpto_src, attendance_dpto)
    shutil.copy2(dpto_src, enrollment_dpto)
    
    # Process attendance
    for path in [attendance_raimon, attendance_dpto]:
        doc = docx.Document(path)
        # Remove the table (table[0])
        if doc.tables:
            table = doc.tables[0]
            tbl_element = table._element
            tbl_element.getparent().remove(tbl_element)
        
        # Replace the paragraph of notes
        replaced = False
        for p in doc.paragraphs:
            if "ha obtenido las calificaciones" in p.text or "calificaciones siguientes" in p.text:
                p.text = "CERTIFICA: Que el alumno <<NombreAlumno>> con documento <<DocumentoIdentidad>> ha asistido regularmente a las clases del programa académico <<nombreCurso>> en el periodo <<añoCurso>>."
                replaced = True
                break
        
        if not replaced:
            for p in doc.paragraphs:
                if "calificaciones" in p.text:
                    p.text = "CERTIFICA: Que el alumno <<NombreAlumno>> con documento <<DocumentoIdentidad>> ha asistido regularmente a las clases del programa académico <<nombreCurso>> en el periodo <<añoCurso>>."
                    replaced = True
                    break
        
        doc.save(path)
        print(f"Generated {os.path.basename(path)}")
        
    # Process enrollment
    for path in [enrollment_raimon, enrollment_dpto]:
        doc = docx.Document(path)
        # Remove the table (table[0])
        if doc.tables:
            table = doc.tables[0]
            tbl_element = table._element
            tbl_element.getparent().remove(tbl_element)
            
        # Replace the paragraph of notes
        replaced = False
        for p in doc.paragraphs:
            if "ha obtenido las calificaciones" in p.text or "calificaciones siguientes" in p.text:
                p.text = "CERTIFICA: Que el alumno <<NombreAlumno>> con documento <<DocumentoIdentidad>> se encuentra matriculado y cursando actualmente el programa académico <<nombreCurso>> en el periodo <<añoCurso>>."
                replaced = True
                break
                
        if not replaced:
            for p in doc.paragraphs:
                if "calificaciones" in p.text:
                    p.text = "CERTIFICA: Que el alumno <<NombreAlumno>> con documento <<DocumentoIdentidad>> se encuentra matriculado y cursando actualmente el programa académico <<nombreCurso>> en el periodo <<añoCurso>>."
                    replaced = True
                    break
                    
        doc.save(path)
        print(f"Generated {os.path.basename(path)}")

if __name__ == "__main__":
    main()
