# -*- coding: utf-8 -*-
import os
import hashlib
from zipfile import ZipFile
from lxml import etree

from docx import Document as DocxDocument

from odoo import fields
from odoo.tests.common import TransactionCase
from odoo.exceptions import ValidationError, UserError


class TestIrgCertificateRequest(TransactionCase):
    """Basic integration tests for irg.certificate.request."""

    def _assert_bottom_right_arcs_are_visible_in_xml(self, res_file):
        with ZipFile(res_file) as docx_zip:
            document_xml = docx_zip.read('word/document.xml').decode(
                'utf-8', errors='ignore'
            )
            rels_xml = docx_zip.read('word/_rels/document.xml.rels').decode(
                'utf-8', errors='ignore'
            )
            package_names = set(docx_zip.namelist())

        self.assertIn('name="Bottom Right Arcs"', document_xml)
        self.assertIn('behindDoc="1"', document_xml)
        self.assertIn('relativeFrom="page"><wp:align>right</wp:align>', document_xml)
        self.assertIn('relativeFrom="page"><wp:align>bottom</wp:align>', document_xml)
        self.assertIn('Target="media/bottom_right_arcs.png"', rels_xml)
        self.assertIn('word/media/bottom_right_arcs.png', package_names)

    def _assert_vertical_legal_text_is_visible_in_xml(self, res_file):
        with ZipFile(res_file) as docx_zip:
            document_xml = ''.join(
                docx_zip.read(name).decode('utf-8', errors='ignore')
                for name in docx_zip.namelist()
                if name.endswith('.xml')
            )

        self.assertIn('B56488687', document_xml)
        self.assertIn('B-603323', document_xml)
        self.assertIn('w:val="10"', document_xml)

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        # Create minimal linked records
        cls.partner = cls.env['res.partner'].create({'name': 'Test Student Portal'})
        cls.course = cls.env['op.course'].create({'name': 'Test Course', 'code': 'TC01'})
        cls.batch = cls.env['op.batch'].create({
            'name': 'Batch A',
            'code': 'BA',
            'course_id': cls.course.id,
            'start_date': fields.Date.today(),
            'end_date': fields.Date.today(),
        })
        cls.product = cls.env['product.product'].create({
            'name': 'Test Course Product',
            'type': 'service',
        })
        cls.register = cls.env['op.admission.register'].create({
            'name': 'Test Register',
            'course_id': cls.course.id,
            'start_date': fields.Date.today(),
            'end_date': fields.Date.today(),
            'min_count': 1,
            'max_count': 100,
            'product_id': cls.product.id,
        })
        cls.admission = cls.env['op.admission'].create({
            'name': 'ADM-TEST',
            'partner_id': cls.partner.id,
            'course_id': cls.course.id,
            'register_id': cls.register.id,
            'gender': 'm',
            'first_name': 'Test',
            'last_name': 'Student',
            'email': 'test.certificate@example.com',
            'birth_date': '1990-01-01',
        })
        # Gradebook template config
        cls.gradebook_tmpl = cls.env['app.gradebook'].create({
            'name': 'Gradebook Template Test',
            'gradebook_template_ids': [(0, 0, {
                'type': 'exam',
                'qty': 2,
                'weight': 100,
            })]
        })
        # Link template to course
        cls.course.write({'gradebook_id': cls.gradebook_tmpl.id})

        cls.gradebook = cls.env['app.gradebook.student'].create({
            'partner_id': cls.partner.id,
            'course_id': cls.course.id,
            'batch_id': cls.batch.id,
            'admission_id': cls.admission.id,
        })

        # Subjects
        cls.subject_normal = cls.env['op.subject'].create({
            'name': 'Subject Compulsory A',
            'code': 'SCA',
            'course_id': cls.course.id,
            'subject_type': 'compulsory',
        })
        cls.subject_pending = cls.env['op.subject'].create({
            'name': 'Subject Compulsory B',
            'code': 'SCB',
            'course_id': cls.course.id,
            'subject_type': 'compulsory',
        })

        # Link subjects to student gradebook
        cls.gb_subj_a = cls.env['app.gradebook.subject'].create({
            'gradebook_student_id': cls.gradebook.id,
            'op_subject_id': cls.subject_normal.id,
        })
        cls.gb_subj_b = cls.env['app.gradebook.subject'].create({
            'gradebook_student_id': cls.gradebook.id,
            'op_subject_id': cls.subject_pending.id,
        })

        # Create results for Subject A
        cls.env['app.gradebook.result'].create({
            'gradebook_subject_id': cls.gb_subj_a.id,
            'survey_type': 'exam',
            'scoring_total': 8.5,
        })
        cls.env['app.gradebook.result'].create({
            'gradebook_subject_id': cls.gb_subj_a.id,
            'survey_type': 'exam',
            'scoring_total': 9.5,
        })
        cls.gb_subj_a.compute_final_subject_note()

        # Create results for Subject B
        cls.env['app.gradebook.result'].create({
            'gradebook_subject_id': cls.gb_subj_b.id,
            'survey_type': 'exam',
            'scoring_total': 7.0,
        })
        cls.gb_subj_b.compute_final_subject_note()

    # ------------------------------------------------------------------
    # Sequence
    # ------------------------------------------------------------------

    def test_01_sequence_assigned_on_create(self):
        """name should be populated from ir.sequence, not stay 'New'."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'digital',
            'state': 'draft',
        })
        self.assertNotEqual(cert.name, 'New', 'Sequence was not assigned on create.')
        self.assertTrue(cert.name.startswith('CERT/'), 'Sequence prefix mismatch.')

    # ------------------------------------------------------------------
    # Price computation
    # ------------------------------------------------------------------

    def test_02_digital_price_no_shipping(self):
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'digital',
            'state': 'draft',
        })
        self.assertAlmostEqual(cert.price_base, 30.0)
        self.assertAlmostEqual(cert.price_shipping, 0.0)
        self.assertAlmostEqual(cert.price_total, 30.0)

    def test_03_physical_national_price(self):
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'physical',
            'shipping_type': 'national',
            'state': 'draft',
        })
        self.assertAlmostEqual(cert.price_base, 40.0)
        self.assertAlmostEqual(cert.price_shipping, 20.0)
        self.assertAlmostEqual(cert.price_total, 60.0)

    # ------------------------------------------------------------------
    # Constraints
    # ------------------------------------------------------------------

    def test_04_physical_without_shipping_raises(self):
        with self.assertRaises(ValidationError):
            self.env['irg.certificate.request'].create({
                'gradebook_student_id': self.gradebook.id,
                'certificate_type': 'physical',
                # shipping_type intentionally omitted
                'state': 'draft',
            })

    # ------------------------------------------------------------------
    # State transitions
    # ------------------------------------------------------------------

    def test_05_cancel_draft(self):
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'digital',
            'state': 'draft',
        })
        cert.action_cancel()
        self.assertEqual(cert.state, 'cancelled')

    def test_06_cannot_cancel_done(self):
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'digital',
            'state': 'done',
        })
        with self.assertRaises(UserError):
            cert.action_cancel()

    def test_07_process_payment_physical(self):
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'physical',
            'shipping_type': 'national',
            'state': 'pending_payment',
        })
        cert._process_payment()
        self.assertEqual(cert.state, 'paid', 'Physical cert should be in "paid" after payment.')

    # ------------------------------------------------------------------
    # app.gradebook.student inherit
    # ------------------------------------------------------------------

    def test_08_certificate_count(self):
        initial_count = self.gradebook.certificate_count
        self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'digital',
            'state': 'done',
        })
        self.gradebook.invalidate_recordset(['certificate_count'])
        self.assertEqual(self.gradebook.certificate_count, initial_count + 1)

    def test_08b_backend_wizard_selection_labels_are_simplified(self):
        fields_info = self.env['irg.certificate.wizard'].fields_get([
            'document_type',
            'certificate_type',
        ])
        self.assertEqual(
            fields_info['document_type']['selection'],
            [
                ('gradebook', 'Certificado de Notas Completo'),
                ('gradebook_partial', 'Certificado de Notas Parcial'),
            ],
        )
        self.assertEqual(
            fields_info['certificate_type']['selection'],
            [
                ('digital', 'Digital'),
                ('physical', 'Físico'),
                ('custom', 'A Medida'),
                ('physical_apostilled', 'Físico Apostillado'),
            ],
        )

    def test_09_bottom_right_arcs_are_added_to_base_word_certificates(self):
        for document_type in ('gradebook', 'attendance', 'enrollment'):
            cert = self.env['irg.certificate.request'].create({
                'gradebook_student_id': self.gradebook.id,
                'certificate_type': 'digital',
                'document_type': document_type,
                'state': 'draft',
            })
            res_file = cert._fill_template()
            self._assert_bottom_right_arcs_are_visible_in_xml(res_file)

    def test_10_final_gradebook_layout_matches_partial_structure(self):
        for certificate_type, signer, signature_lines in (
            ('digital', 'raimon', ['Raimon Gaja Jaumeandreu', 'Instituto Raimon Gaja']),
            ('digital', 'dpto_academico', ['Departamento Académico', 'Instituto Raimon Gaja']),
            ('physical', 'raimon', ['Raimon Gaja Jaumeandreu', 'Instituto Raimon Gaja']),
        ):
            cert = self.env['irg.certificate.request'].create({
                'gradebook_student_id': self.gradebook.id,
                'certificate_type': certificate_type,
                'shipping_type': 'national' if certificate_type == 'physical' else False,
                'document_type': 'gradebook',
                'signer': signer,
                'state': 'draft',
            })

            res_file = cert._fill_template()
            self._assert_vertical_legal_text_is_visible_in_xml(res_file)
            document = DocxDocument(res_file)

            certifica = next(
                (para for para in document.paragraphs if para.text.strip() == 'CERTIFICA:'),
                None,
            )
            closing = next(
                (
                    para for para in document.paragraphs
                    if 'Para que así conste' in para.text
                ),
                None,
            )
            signature_paragraph = next(
                (
                    para for para in document.paragraphs
                    if [line.strip() for line in para.text.splitlines()] == signature_lines
                ),
                None,
            )
            first_body = next(
                (
                    para for para in document.paragraphs
                    if para.text.startswith('Que Test Student Portal con')
                ),
                None,
            )
            second_body = next(
                (
                    para for para in document.paragraphs
                    if para.text.startswith('Que, el Máster consta de')
                ),
                None,
            )
            third_body = next(
                (
                    para for para in document.paragraphs
                    if para.text == 'Las calificaciones obtenidas son:'
                ),
                None,
            )

            self.assertIsNotNone(certifica)
            self.assertEqual(certifica.alignment, 0)
            self.assertEqual(certifica.paragraph_format.left_indent.twips, -172)
            self.assertEqual(certifica.paragraph_format.right_indent.twips, -783)
            self.assertIsNotNone(first_body)
            self.assertIn(' ha realizado y superado el ', first_body.text)
            self.assertNotIn('consta matriculado', first_body.text)
            self.assertIn('durante el período académico', first_body.text)
            self.assertNotIn('ha obtenido las calificaciones siguientes', first_body.text)
            self.assertEqual(first_body.alignment, 3)
            self.assertEqual(first_body.paragraph_format.left_indent.twips, -172)
            self.assertEqual(first_body.paragraph_format.right_indent.twips, -783)
            bold_runs = [run.text for run in first_body.runs if run.bold]
            self.assertIn('Test Student Portal', bold_runs)
            self.assertIn('Test Course', bold_runs)
            self.assertIsNotNone(second_body)
            self.assertIn(
                '60 ECTS, equivalentes a 1500 horas de estudio',
                second_body.text,
            )
            self.assertEqual(second_body.alignment, 3)
            self.assertEqual(second_body.paragraph_format.left_indent.twips, -172)
            self.assertEqual(second_body.paragraph_format.right_indent.twips, -783)
            self.assertIsNotNone(third_body)
            self.assertEqual(third_body.alignment, 3)
            self.assertEqual(third_body.paragraph_format.left_indent.twips, -172)
            self.assertEqual(third_body.paragraph_format.right_indent.twips, -783)
            self.assertIsNotNone(closing)
            self.assertEqual(closing.alignment, 3)
            self.assertEqual(closing.paragraph_format.left_indent.twips, -172)
            self.assertEqual(closing.paragraph_format.right_indent.twips, -783)
            self.assertIsNotNone(signature_paragraph)
            self.assertEqual([line.strip() for line in signature_paragraph.text.splitlines()], signature_lines)
            self.assertEqual(signature_paragraph.alignment, 0)
            self.assertEqual(signature_paragraph.paragraph_format.left_indent.twips, -172)
            self.assertEqual(signature_paragraph.paragraph_format.right_indent.twips, -783)

    def _assert_bottom_right_arcs_are_absent_in_xml(self, res_file):
        with ZipFile(res_file) as docx_zip:
            document_xml = docx_zip.read('word/document.xml').decode(
                'utf-8', errors='ignore'
            )
            rels_xml = docx_zip.read('word/_rels/document.xml.rels').decode(
                'utf-8', errors='ignore'
            )
            package_names = set(docx_zip.namelist())

        self.assertNotIn('name="Bottom Right Arcs"', document_xml)
        self.assertNotIn('Target="media/bottom_right_arcs.png"', rels_xml)
        self.assertNotIn('word/media/bottom_right_arcs.png', package_names)

    def _assert_header_logo_is_removed_in_xml(self, res_file):
        header_name = 'word/header1.xml'
        with ZipFile(res_file) as docx_zip:
            if header_name not in docx_zip.namelist():
                return
            header_xml = etree.fromstring(docx_zip.read(header_name))
        
        runs_with_drawings = []
        for r in header_xml.xpath('.//*[local-name()="r"]'):
            if r.xpath('.//*[local-name()="drawing" or local-name()="AlternateContent" or local-name()="pict"]'):
                runs_with_drawings.append(r)
        self.assertEqual(len(runs_with_drawings), 0, "Header logo drawings were not removed.")

    def _assert_header_logo_is_present_in_xml(self, res_file):
        header_name = 'word/header1.xml'
        with ZipFile(res_file) as docx_zip:
            if header_name not in docx_zip.namelist():
                self.fail("Header XML not found in document.")
            header_xml = etree.fromstring(docx_zip.read(header_name))
        
        runs_with_drawings = []
        for r in header_xml.xpath('.//*[local-name()="r"]'):
            if r.xpath('.//*[local-name()="drawing" or local-name()="AlternateContent" or local-name()="pict"]'):
                runs_with_drawings.append(r)
        self.assertTrue(len(runs_with_drawings) > 0, "Header logo drawings were removed but should be present.")

    def test_11_physical_gradebook_omits_logo_and_arcs(self):
        """Physical/apostilled gradebook certificates must omit the header logo and decorative arcs."""
        for cert_type in ('physical', 'physical_apostilled'):
            cert = self.env['irg.certificate.request'].create({
                'gradebook_student_id': self.gradebook.id,
                'document_type': 'gradebook',
                'certificate_type': cert_type,
                'shipping_type': 'national',
                'state': 'draft',
            })
            res_file = cert._fill_template()
            self._assert_bottom_right_arcs_are_absent_in_xml(res_file)
            self._assert_header_logo_is_removed_in_xml(res_file)

    def test_12_non_physical_gradebook_retains_logo_and_arcs(self):
        """Digital/custom gradebook certificates must retain the header logo and decorative arcs."""
        for cert_type in ('digital', 'custom'):
            cert = self.env['irg.certificate.request'].create({
                'gradebook_student_id': self.gradebook.id,
                'document_type': 'gradebook',
                'certificate_type': cert_type,
                'state': 'draft',
            })
            res_file = cert._fill_template()
            self._assert_bottom_right_arcs_are_visible_in_xml(res_file)
            self._assert_header_logo_is_present_in_xml(res_file)

    def _assert_table_font_sizes_match_top_font_size(self, res_file):
        doc = DocxDocument(res_file)
        top_font_size = None
        for para in doc.paragraphs:
            if para.text.strip():
                for r in para.runs:
                    if r.font and r.font.size:
                        top_font_size = r.font.size
                        break
            if top_font_size:
                break
        self.assertIsNotNone(top_font_size, "Could not find top font size in document.")
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            self.assertEqual(
                                run.font.size,
                                top_font_size,
                                "Table run font size does not match calculated top font size."
                            )

    def _assert_signature_logo_is_present_and_referenced(self, res_file):
        module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        logo_path = os.path.join(module_path, 'static', 'src', 'img', 'logodesgastado.png')
        self.assertTrue(os.path.isfile(logo_path), "Expected signature logo does not exist in module path.")
        with open(logo_path, 'rb') as f:
            expected_md5 = hashlib.md5(f.read()).hexdigest()

        with ZipFile(res_file) as z:
            matching_media_file = None
            for name in z.namelist():
                if name.startswith('word/media/'):
                    data = z.read(name)
                    if hashlib.md5(data).hexdigest() == expected_md5:
                        matching_media_file = name
                        break
            self.assertIsNotNone(matching_media_file, "logodesgastado.png not found by MD5 hash inside ZIP media.")
            
            doc_xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
            rels_xml = z.read('word/_rels/document.xml.rels').decode('utf-8', errors='ignore')
            
        relative_target = matching_media_file.replace('word/', '')
        import re
        r_ids = re.findall(rf'Id="([^"]+)"[^>]+Target="{re.escape(relative_target)}"', rels_xml)
        self.assertTrue(len(r_ids) > 0, f"Relationship for {relative_target} not found in document.xml.rels")
        
        referenced = any(r_id in doc_xml for r_id in r_ids)
        self.assertTrue(referenced, f"Relationship ID(s) {r_ids} for signature logo not referenced in document.xml")

        # Validate that the signature logo is placed in the same paragraph containing the handwritten signature, having at least 2 drawing elements
        doc = DocxDocument(res_file)
        sig_rel_id = None
        for rel_id, rel in doc.part.rels.items():
            target = getattr(rel, 'target_ref', '').lower()
            if any(img in target for img in ('media/image2.jpg', 'media/image2.png', 'media/image2.jpeg')):
                sig_rel_id = rel_id
                break
        self.assertIsNotNone(sig_rel_id, "Handwritten signature rel_id not found")

        paragraphs_to_search = list(doc.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    paragraphs_to_search.extend(cell.paragraphs)

        sig_para = None
        for para in paragraphs_to_search:
            for r in para.runs:
                embed_nodes = r._r.xpath('.//*[@*[local-name()="embed" and .="%s"]]' % sig_rel_id)
                if embed_nodes:
                    sig_para = para
                    break
            if sig_para:
                break

        self.assertIsNotNone(sig_para, "Paragraph containing the handwritten signature not found")
        drawing_runs = [r for r in sig_para.runs if r._r.xpath('.//*[local-name()="drawing" or local-name()="pict"]')]
        self.assertTrue(len(drawing_runs) >= 2, "Paragraph containing the handwritten signature does not contain at least 2 drawings")

    def _assert_table_font_sizes_match_value(self, res_file, expected_size):
        doc = DocxDocument(res_file)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            self.assertEqual(
                                run.font.size,
                                expected_size,
                                f"Table run font size is {run.font.size}, expected {expected_size}."
                            )

    def test_13_table_font_sizes_match_top_font_size(self):
        """Verify that table cell font size matches the top font size for digital and is 7.5 Pt for physical."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'digital',
            'document_type': 'gradebook',
            'state': 'draft',
        })
        res_file = cert._fill_template()
        self._assert_table_font_sizes_match_top_font_size(res_file)

        cert_phys = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'physical',
            'shipping_type': 'national',
            'document_type': 'gradebook',
            'state': 'draft',
        })
        res_file_phys = cert_phys._fill_template()
        self._assert_table_font_sizes_match_value(res_file_phys, Pt(7.5))

    def test_14_signature_logo_present_for_raimon_signer(self):
        """Verify that Raimon Gaja's signature logo is present in the ZIP and referenced in document.xml."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'digital',
            'document_type': 'gradebook',
            'signer': 'raimon',
            'state': 'draft',
        })
        res_file = cert._fill_template()
        self._assert_signature_logo_is_present_and_referenced(res_file)

    def _assert_table_data_row_heights_are_315_atleast(self, res_file):
        from docx.oxml.ns import qn
        doc = DocxDocument(res_file)
        table = doc.tables[0]
        rows = list(table.rows)
        self.assertTrue(len(rows) > 2, "Table should have header, data, and footer rows")
        
        data_rows = rows[1:-1]
        for idx, row in enumerate(data_rows):
            trPr = row._tr.find(qn('w:trPr'))
            self.assertIsNotNone(trPr, f"Row {idx+1} is missing trPr")
            trHeight = trPr.find(qn('w:trHeight'))
            self.assertIsNotNone(trHeight, f"Row {idx+1} is missing trHeight")
            
            val = trHeight.get(qn('w:val'))
            hRule = trHeight.get(qn('w:hRule'))
            
            self.assertEqual(val, '315', f"Row {idx+1} height val is {val}, expected '315'")
            self.assertEqual(hRule, 'atLeast', f"Row {idx+1} height hRule is {hRule}, expected 'atLeast'")

    def test_15_table_data_row_heights_are_315_atleast(self):
        """Verify that all data rows in the grades table have height=315 dxa and hRule=atLeast."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'digital',
            'document_type': 'gradebook',
            'state': 'draft',
        })
        res_file = cert._fill_template()
        self._assert_table_data_row_heights_are_315_atleast(res_file)

        cert_phys = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'physical',
            'shipping_type': 'national',
            'document_type': 'gradebook',
            'state': 'draft',
        })
        res_file_phys = cert_phys._fill_template()
        self._assert_table_data_row_heights_are_315_atleast(res_file_phys)

    def test_16_gradebook_dpto_academico_intro_replacement(self):
        """Verify that when using 'dpto_academico' signer, the introduction paragraph is replaced correctly."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'digital',
            'document_type': 'gradebook',
            'signer': 'dpto_academico',
            'state': 'draft',
        })
        res_file = cert._fill_template()
        document = DocxDocument(res_file)
        intro = next(
            (
                para for para in document.paragraphs
                if 'El Instituto Raimon Gaja, con CIF B-56488687' in para.text
            ),
            None,
        )
        self.assertIsNotNone(intro, "Introduction paragraph not found for dpto_academico signer.")
        self.assertEqual(
            intro.text.strip(),
            'El Instituto Raimon Gaja, con CIF B-56488687 en calle Córcega 213, 1º 2ª, 08036 Barcelona.'
        )

    def test_17_gradebook_course_id_4_ects_text(self):
        """Verify ECTS text and detailed text when course ID is 4 for gradebook."""
        course_4 = self.env['op.course'].browse(4)
        if not course_4.exists():
            course_tmp = self.env['op.course'].create({
                'name': 'Special Course 4',
                'code': 'SC04',
            })
            self.env.cr.execute("UPDATE op_course SET id = 4 WHERE id = %s", (course_tmp.id,))
            self.env.registry.clear_cache()
            course_4 = self.env['op.course'].browse(4)
        
        course_4.write({'gradebook_id': self.gradebook_tmpl.id})
        
        batch_4 = self.env['op.batch'].search([('course_id', '=', course_4.id)], limit=1)
        if not batch_4:
            batch_4 = self.env['op.batch'].create({
                'name': 'Batch 4',
                'code': 'B4',
                'course_id': course_4.id,
                'start_date': fields.Date.today(),
                'end_date': fields.Date.today(),
            })
        
        register_4 = self.env['op.admission.register'].create({
            'name': 'Register 4',
            'course_id': course_4.id,
            'start_date': fields.Date.today(),
            'end_date': fields.Date.today(),
            'min_count': 1,
            'max_count': 100,
            'product_id': self.product.id,
        })
        
        admission_4 = self.env['op.admission'].create({
            'name': 'ADM-TEST-4',
            'partner_id': self.partner.id,
            'course_id': course_4.id,
            'register_id': register_4.id,
            'gender': 'm',
            'first_name': 'Test4',
            'last_name': 'Student4',
            'email': 'test4@example.com',
            'birth_date': '1990-01-01',
        })
        
        gradebook_4 = self.env['app.gradebook.student'].create({
            'partner_id': self.partner.id,
            'course_id': course_4.id,
            'batch_id': batch_4.id,
            'admission_id': admission_4.id,
        })
        
        self.env['app.gradebook.subject'].create({
            'gradebook_student_id': gradebook_4.id,
            'op_subject_id': self.subject_normal.id,
        })
        
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': gradebook_4.id,
            'certificate_type': 'digital',
            'document_type': 'gradebook',
            'state': 'draft',
        })
        
        res_file = cert._fill_template()
        document = DocxDocument(res_file)
        
        second_body = next(
            (
                para for para in document.paragraphs
                if '120 ECTS' in para.text
            ),
            None,
        )
        self.assertIsNotNone(second_body, "120 ECTS text not found in generated document.")
        self.assertIn('120 ECTS, equivalentes a 3000 horas de estudio', second_body.text)

    def test_18_physical_gradebook_modifications(self):
        """Verify top margin, unscaled outer fonts, table cell fonts, and signature/stamp removal on physical gradebooks."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'certificate_type': 'physical',
            'shipping_type': 'national',
            'document_type': 'gradebook',
            'signer': 'raimon',
            'state': 'draft',
        })
        res_file = cert._fill_template()
        doc = DocxDocument(res_file)
        
        # 1. Verify top margin shift (72 Pt default + 56.25 Pt shift = 128.25 Pt)
        self.assertEqual(doc.sections[0].top_margin.pt, 128.25)
        
        # 2. Verify outer font size is 10 Pt (not scaled down by 75%)
        body_runs = [r for p in doc.paragraphs if p.text.strip() for r in p.runs if r.font and r.font.size]
        self.assertTrue(len(body_runs) > 0, "No body runs found with font size set.")
        for r in body_runs:
            self.assertEqual(r.font.size.pt, 10.0, f"Outer text run '{r.text}' font size is {r.font.size.pt}, expected 10.0 Pt")
            
        # 3. Verify table run font size is 7.5 Pt (95250 EMUs)
        self._assert_table_font_sizes_match_value(res_file, Pt(7.5))
        
        # 4. Verify signature/stamp images are removed (media/image2.jpg / media/image2.png / media/image2.jpeg)
        sig_rel_ids = []
        for rel_id, rel in doc.part.rels.items():
            target = getattr(rel, 'target_ref', '').lower()
            if any(img in target for img in ('media/image2.jpg', 'media/image2.png', 'media/image2.jpeg')):
                sig_rel_ids.append(rel_id)
        
        embeds = []
        if sig_rel_ids:
            for para in list(doc.paragraphs) + [p for t in doc.tables for r in t.rows for c in r.cells for p in c.paragraphs]:
                for r in para.runs:
                    for rel_id in sig_rel_ids:
                        embeds.extend(r._r.xpath('.//*[@*[local-name()="embed" and .="%s"]]' % rel_id))
        self.assertEqual(len(embeds), 0, "Signature/stamp images still embedded in document XML.")


