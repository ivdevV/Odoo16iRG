# -*- coding: utf-8 -*-
import base64
import logging
import os
import subprocess
import tempfile
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Pt, Twips
from copy import deepcopy
from lxml import etree

from odoo import models, fields, api, _
from odoo.exceptions import UserError, ValidationError

_logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

CERTIFICATE_TYPES = [
    ('digital', 'Certificado de Notas Digital'),
    ('physical', 'Certificado de Notas Físico'),
    ('custom', 'Certificado de Notas a Medida'),
    ('physical_apostilled', 'Certificado de Notas Físico Apostillado'),
]

SHIPPING_TYPES = [
    ('national', 'Envío Nacional'),
    ('international', 'Envío Internacional'),
]

STATE_SELECTION = [
    ('draft', 'Borrador'),
    ('pending_payment', 'Pendiente de Pago'),
    ('paid', 'Pagado'),
    ('in_process', 'En Proceso'),
    ('sent', 'Enviado'),
    ('done', 'Finalizado'),
    ('cancelled', 'Cancelado'),
]

CUSTOM_OPTIONS = [
    ('language_en', 'En Inglés'),
    ('language_fr', 'En Francés'),
    ('specific_subjects', 'Asignaturas Específicas'),
    ('official_seal', 'Con Sello Oficial'),
]

PRICE_MAP = {
    'digital': 30.0,
    'physical': 40.0,
    'custom': 40.0,
    'physical_apostilled': 120.0,
}

SHIPPING_MAP = {
    'national': 20.0,
    'international': 60.0,
}

PHYSICAL_TYPES = ('physical', 'physical_apostilled')

SIGNER_SELECTION = [
    ('raimon', 'Raimon Gaja Jaumeandreu'),
    ('dpto_academico', 'Departamento Académico'),
]

# Product xml ids used when generating portal invoices
_PORTAL_PRODUCT_XMLID = {
    'digital': 'irg_gradebook_certificates.product_cert_digital',
    'physical': 'irg_gradebook_certificates.product_cert_physical',
    'custom': 'irg_gradebook_certificates.product_cert_custom',
    'physical_apostilled': 'irg_gradebook_certificates.product_cert_apostilled',
}

_PORTAL_SHIPPING_XMLID = {
    'national': 'irg_gradebook_certificates.product_shipping_national',
    'international': 'irg_gradebook_certificates.product_shipping_international',
}

# Keyword that identifies the MNC course (Máster en Neuropsicología Clínica
# basada en la Evidencia) to apply its specific ECTS / duration values.
_MNC_KEYWORD = 'Neuropsicología'


class IrgCertificateRequest(models.Model):
    _name = 'irg.certificate.request'
    _description = 'Solicitud de Certificado de Notas'
    _inherit = ['mail.thread', 'mail.activity.mixin']
    _order = 'id desc'

    _GRADEBOOK_TEXT_INDENT = Twips(-172)
    _GRADEBOOK_TABLE_WIDTH = Twips(9010)
    _GRADEBOOK_PAGE_TEXT_WIDTH = Twips(8055)
    _GRADEBOOK_TEXT_RIGHT_INDENT = (
        _GRADEBOOK_PAGE_TEXT_WIDTH - _GRADEBOOK_TEXT_INDENT - _GRADEBOOK_TABLE_WIDTH
    )
    _GRADEBOOK_COMPACT_TOP_MARGIN = Pt(48)
    _GRADEBOOK_COMPACT_BOTTOM_MARGIN = Pt(20)
    # Spacer under the header logo for many-subject certificates (MNC).
    _GRADEBOOK_COMPACT_LOGO_SPACER_AFTER = Pt(36)
    _DPTO_ACADEMICO_INTRO = 'El Instituto Raimon Gaja, con CIF B-56488687 en calle Córcega 213, 1º 2ª, 08036 Barcelona.'

    def _replace_dpto_academico_intro(self, doc):
        """Replace the default Dpto Académico intro in document paragraphs."""
        if self.signer != 'dpto_academico':
            return
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == 'El Departamento Académico del Instituto Raimon Gaja, S.L.':
                self._replace_in_paragraph(paragraph, 'El Departamento Académico del Instituto Raimon Gaja, S.L.', self._DPTO_ACADEMICO_INTRO)
                self._format_gradebook_body_paragraph(paragraph, justify=True)
                break

    # ------------------------------------------------------------------
    # Identity
    # ------------------------------------------------------------------

    name = fields.Char(
        string='Referencia',
        readonly=True,
        copy=False,
        default='New',
    )
    origin = fields.Selection(
        [('backend', 'Generación Interna'), ('portal', 'Solicitud Portal')],
        string='Origen',
        default='backend',
        readonly=True,
        required=True,
    )
    document_type = fields.Selection(
        selection=[
            ('gradebook', 'Certificado de Notas Completo'),
            ('gradebook_partial', 'Certificado de Notas Parcial'),
            ('diploma', 'Diploma'),
            ('attendance', 'Certificado de Asistencia'),
            ('enrollment', 'Certificado de Matrícula'),
        ],
        string='Tipo de Documento',
        default='gradebook',
        required=True,
        tracking=True,
    )

    # ------------------------------------------------------------------
    # Academic data (from gradebook)
    # ------------------------------------------------------------------

    gradebook_student_id = fields.Many2one(
        'app.gradebook.student',
        string='Libreta',
        required=True,
        ondelete='restrict',
        tracking=True,
    )
    partner_id = fields.Many2one(
        'res.partner',
        string='Alumno',
        related='gradebook_student_id.partner_id',
        store=True,
    )
    admission_id = fields.Many2one(
        'op.admission',
        string='Admisión',
        related='gradebook_student_id.admission_id',
        store=True,
    )
    course_id = fields.Many2one(
        'op.course',
        string='Programa',
        related='gradebook_student_id.course_id',
        store=True,
    )
    batch_id = fields.Many2one(
        'op.batch',
        string='Grupo',
        related='gradebook_student_id.batch_id',
        store=True,
    )

    # ------------------------------------------------------------------
    # Certificate options
    # ------------------------------------------------------------------

    certificate_type = fields.Selection(
        selection=CERTIFICATE_TYPES,
        string='Tipo de Certificado',
        required=True,
        tracking=True,
    )
    shipping_type = fields.Selection(
        selection=SHIPPING_TYPES,
        string='Tipo de Envío',
        tracking=True,
    )
    custom_description = fields.Text(
        string='Descripción de Peticiones',
    )
    custom_options = fields.Selection(
        selection=CUSTOM_OPTIONS,
        string='Opción Adicional',
    )
    signer = fields.Selection(
        selection=SIGNER_SELECTION,
        string='Persona que Firma',
        default='raimon',
        tracking=True,
    )

    # ------------------------------------------------------------------
    # State / lifecycle
    # ------------------------------------------------------------------

    state = fields.Selection(
        selection=STATE_SELECTION,
        string='Estado',
        default='draft',
        required=True,
        tracking=True,
        copy=False,
    )
    request_date = fields.Datetime(
        string='Fecha de Solicitud',
        default=fields.Datetime.now,
        readonly=True,
        copy=False,
    )
    delivery_date = fields.Datetime(
        string='Fecha de Entrega/Envío',
        copy=False,
        tracking=True,
    )
    tracking_number = fields.Char(
        string='Número de Seguimiento',
        copy=False,
        tracking=True,
    )
    apostille_done = fields.Boolean(
        string='Apostillado Tramitado',
        default=False,
        tracking=True,
    )

    # ------------------------------------------------------------------
    # Pricing
    # ------------------------------------------------------------------

    price_base = fields.Float(
        string='Precio Base (€)',
        compute='_compute_prices',
        store=True,
        digits=(6, 2),
    )
    price_shipping = fields.Float(
        string='Precio Envío (€)',
        compute='_compute_prices',
        store=True,
        digits=(6, 2),
    )
    price_total = fields.Float(
        string='Total (€)',
        compute='_compute_prices',
        store=True,
        digits=(6, 2),
    )

    # ------------------------------------------------------------------
    # Linked records
    # ------------------------------------------------------------------

    sale_order_id = fields.Many2one(
        'sale.order',
        string='Pedido de Venta',
        copy=False,
        readonly=True,
    )
    invoice_id = fields.Many2one(
        'account.move',
        string='Factura de Portal',
        copy=False,
        readonly=True,
    )
    payment_link = fields.Char(
        string='Enlace de Pago',
        copy=False,
        readonly=True,
    )
    attachment_id = fields.Many2one(
        'ir.attachment',
        string='PDF Generado',
        copy=False,
        readonly=True,
    )
    internal_notes = fields.Text(
        string='Notas Internas',
    )

    # ------------------------------------------------------------------
    # ORM overrides
    # ------------------------------------------------------------------

    @api.model
    def create(self, vals):
        if vals.get('name', 'New') == 'New':
            vals['name'] = (
                self.env['ir.sequence'].next_by_code('irg.certificate.request') or 'New'
            )
        return super().create(vals)

    # ------------------------------------------------------------------
    # Computed
    # ------------------------------------------------------------------

    @api.depends('certificate_type', 'shipping_type')
    def _compute_prices(self):
        for rec in self:
            base = PRICE_MAP.get(rec.certificate_type, 0.0)
            shipping = (
                SHIPPING_MAP.get(rec.shipping_type, 0.0)
                if rec.certificate_type in PHYSICAL_TYPES
                else 0.0
            )
            rec.price_base = base
            rec.price_shipping = shipping
            rec.price_total = base + shipping

    # ------------------------------------------------------------------
    # Constraints
    # ------------------------------------------------------------------

    @api.constrains('certificate_type', 'shipping_type')
    def _check_shipping_required(self):
        for rec in self:
            if rec.certificate_type in PHYSICAL_TYPES and not rec.shipping_type:
                raise ValidationError(
                    _('El tipo de envío es obligatorio para certificados físicos.')
                )

    # ------------------------------------------------------------------
    # State transitions (backend buttons)
    # ------------------------------------------------------------------

    def action_cancel(self):
        for rec in self:
            if rec.state in ('done', 'sent'):
                raise UserError(
                    _(
                        'No se puede cancelar el certificado "%s": '
                        'ya está en estado "%s".',
                        rec.name,
                        dict(STATE_SELECTION).get(rec.state),
                    )
                )
            rec.state = 'cancelled'

    def action_mark_in_process(self):
        for rec in self:
            if rec.state != 'paid':
                raise UserError(_('Solo se puede procesar un certificado pagado.'))
            rec.state = 'in_process'

    def action_mark_sent(self):
        """Wizard-less version: opens a simple wizard to enter tracking number."""
        self.ensure_one()
        return {
            'type': 'ir.actions.act_window',
            'name': _('Marcar como Enviado'),
            'res_model': 'irg.certificate.sent.wizard',
            'view_mode': 'form',
            'target': 'new',
            'context': {'default_certificate_id': self.id},
        }

    def _do_mark_sent(self, tracking_number=None):
        """Called from irg.certificate.sent.wizard."""
        for rec in self:
            if rec.state != 'in_process':
                raise UserError(
                    _('El certificado debe estar "En Proceso" para marcarse como enviado.')
                )
            rec.write({
                'state': 'sent',
                'delivery_date': fields.Datetime.now(),
                'tracking_number': tracking_number or rec.tracking_number,
            })
            rec._send_sent_notification()

    def action_mark_done(self):
        for rec in self:
            rec.state = 'done'

    # ------------------------------------------------------------------
    # Payment processing
    # ------------------------------------------------------------------

    def _process_payment(self):
        """
        Called by sale.order.action_confirm() hook after payment.
        Digital/custom: generate PDF and mark done.
        Physical: mark paid and queue for admin processing.
        """
        self.ensure_one()
        self.state = 'paid'
        if self.certificate_type in ('digital', 'custom'):
            self._generate_and_attach_pdf()
            self.state = 'done'
            self._send_digital_notification()
        else:
            self._send_paid_notification()
            # Notify academic team so they can prepare and ship the physical cert
            self._send_team_notification()

    # ------------------------------------------------------------------
    # PDF generation
    # ------------------------------------------------------------------

    # ------------------------------------------------------------------
    # Docx template helpers
    # ------------------------------------------------------------------

    def _get_template_path(self):
        """Return the absolute path to the Word certificate template based on signer and document_type."""
        module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        tpl_dir = os.path.join(module_path, 'static', 'src', 'templates')
        
        signer_suffix = 'dpto' if self.signer == 'dpto_academico' else 'raimon'
        
        if self.document_type == 'attendance':
            filename = f'Plantilla-certificado-asistencia-{signer_suffix}.docx'
        elif self.document_type == 'enrollment':
            filename = f'Plantilla-certificado-curso-{signer_suffix}.docx'
        else:
            # Fallback for gradebook / gradebook_partial
            filename = f'Plantilla-certificado-notas-{signer_suffix}.docx'
            
        return os.path.join(tpl_dir, filename)

    @staticmethod
    def _replace_in_paragraph(paragraph, old, new):
        """Replace *old* with *new* across runs that Word may have split."""
        full = ''.join(r.text for r in paragraph.runs)
        if old not in full:
            return
        full = full.replace(old, new)
        if paragraph.runs:
            paragraph.runs[0].text = full
            for r in paragraph.runs[1:]:
                for t in r._element.xpath('.//w:t'):
                    t.text = ''

    @staticmethod
    def _replace_paragraph_text_with_bold_segments(paragraph, segments):
        """Replace paragraph text with runs, preserving first-run style."""
        base_run = paragraph.runs[0] if paragraph.runs else None
        base_rpr = None
        if base_run is not None and base_run._r.rPr is not None:
            base_rpr = deepcopy(base_run._r.rPr)

        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)

        for text, bold in segments:
            if not text:
                continue
            run = paragraph.add_run(text)
            if base_rpr is not None:
                run._r.insert(0, deepcopy(base_rpr))
            run.bold = bool(bold)

    @staticmethod
    def _scale_document_fonts(doc, percent=75):
        """Scale font sizes and paragraph spacing in the Word document.

        Covers both document.xml (runs) and styles.xml (named styles/defaults),
        which is where most templates store font definitions.
        Also reduces paragraph before/after spacing to compact the layout.

        Word stores font sizes in half-points: w:val="22" = 11 pt.
        Paragraph spacing is in twentieths of a point (twips).
        Minimum font enforced: 14 half-points (7 pt).
        """
        NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        w_val = '{%s}val' % NS_W
        sz_tags = ('{%s}sz' % NS_W, '{%s}szCs' % NS_W)
        spacing_tag = '{%s}spacing' % NS_W
        w_before = '{%s}before' % NS_W
        w_after = '{%s}after' % NS_W

        # Collect all XML roots: document body + styles (where fonts usually live)
        roots = [doc.element]
        try:
            roots.append(doc.part.styles._element)
        except Exception:
            pass

        for root in roots:
            # Scale explicit font sizes
            for tag in sz_tags:
                for el in root.iter(tag):
                    raw = el.get(w_val)
                    if raw and raw.isdigit():
                        new_val = max(14, int(round(int(raw) * percent / 100)))
                        el.set(w_val, str(new_val))

            # Compact paragraph spacing (before/after)
            for el in root.iter(spacing_tag):
                for attr in (w_before, w_after):
                    raw = el.get(attr)
                    if raw and raw.isdigit():
                        new_val = max(0, int(round(int(raw) * percent / 100)))
                        el.set(attr, str(new_val))

    def _format_gradebook_body_paragraph(self, paragraph, justify=True):
        """Align gradebook text blocks with the notes table grid."""
        fmt = paragraph.paragraph_format
        fmt.left_indent = self._GRADEBOOK_TEXT_INDENT
        fmt.right_indent = self._GRADEBOOK_TEXT_RIGHT_INDENT
        if justify:
            paragraph.alignment = 3  # WD_ALIGN_PARAGRAPH.JUSTIFY
        return paragraph

    def _format_gradebook_signature_paragraph(self, paragraph):
        """Stack signer and institute lines without altering unrelated text."""
        is_physical = self.certificate_type in PHYSICAL_TYPES
        normalized_text = ' '.join(paragraph.text.split())
        if normalized_text == 'Departamento Académico Instituto Raimon Gaja':
            paragraph.text = 'Departamento Académico\nInstituto Raimon Gaja'
        elif normalized_text in (
            'Raimon Gaja Jaumeandreu Instituto Raimon Gaja',
            'Raimon Gaja Instituto Raimon Gaja'
        ):
            if is_physical:
                paragraph.text = 'Raimon Gaja\nDirector General iRG'
            else:
                paragraph.text = 'Raimon Gaja Jaumeandreu\nInstituto Raimon Gaja'
        self._format_gradebook_body_paragraph(paragraph, justify=False)
        paragraph.alignment = 0  # WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.tab_stops.clear_all()
        for run in paragraph.runs:
            if run.text:
                run.text = run.text.lstrip()
            if is_physical:
                run.font.size = Pt(9.25)
        return paragraph

    def _format_gradebook_static_paragraphs(self, doc):
        """Normalize fixed certificate blocks to match the partial layout."""
        is_physical = self.certificate_type in PHYSICAL_TYPES
        signer_intro_markers = (
            'Raimon Gaja Jaumeandreu, con DNI',
            'Raimon Gaja, con DNI',
            self._DPTO_ACADEMICO_INTRO,
        )
        signature_texts = (
            'Raimon Gaja Jaumeandreu Instituto Raimon Gaja',
            'Raimon Gaja Instituto Raimon Gaja',
            'Departamento Académico Instituto Raimon Gaja',
        )
        for para in doc.paragraphs:
            text = para.text.strip()
            normalized_text = ' '.join(text.split())
            if not text:
                continue
            if text == 'CERTIFICA':
                para.text = 'CERTIFICA:'
                text = para.text.strip()
            if (
                text == 'CERTIFICA:'
                or any(marker in text for marker in signer_intro_markers)
            ):
                self._format_gradebook_body_paragraph(para, justify=False)
                para.alignment = 0  # WD_ALIGN_PARAGRAPH.LEFT
            if 'Para que así conste' in text:
                if is_physical:
                    fecha_larga = ''
                    if self.request_date:
                        meses = {
                            1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
                            5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
                            9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre',
                        }
                        dt = self.request_date
                        fecha_larga = '%d de %s de %d' % (dt.day, meses[dt.month], dt.year)
                    para.text = f'Para que así conste, firmo la presente en Barcelona, a fecha {fecha_larga}.'
                    para.paragraph_format.space_after = Pt(48)
                    for run in para.runs:
                        run.font.size = Pt(9.25)
                self._format_gradebook_body_paragraph(para, justify=True)
            if normalized_text in signature_texts:
                self._format_gradebook_signature_paragraph(para)

    def _replace_gradebook_description_paragraph(
        self, doc, partner, documento, course_name, periodo_str, ects_detallado, has_many_subjects=False
    ):
        """Use the partial certificate body wording in the final gradebook."""
        sentence_2 = (
            'Que, el Máster consta de %s, distribuidas entre horas de clases '
            'y horas destinadas a otras actividades académicas.'
        ) % ects_detallado
        sentence_3 = 'Las calificaciones obtenidas son:'

        sp_after_body = Pt(4) if has_many_subjects else Pt(12)
        sp_after_p3 = Pt(7) if has_many_subjects else Pt(12)

        for para in list(doc.paragraphs):
            full_text = ''.join(r.text for r in para.runs)
            if 'ha obtenido las calificaciones siguientes:' not in full_text:
                continue

            self._replace_paragraph_text_with_bold_segments(para, [
                ('Que ', False),
                (partner.name or '', True),
                (' con %s ha realizado y superado el ' % documento, False),
                (course_name, True),
                (' durante el período académico %s.' % periodo_str, False),
            ])
            self._format_gradebook_body_paragraph(para, justify=True)
            para.paragraph_format.space_after = sp_after_body

            p_2 = doc.add_paragraph(sentence_2)
            p_2.style = para.style
            self._format_gradebook_body_paragraph(p_2, justify=True)
            p_2.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
            p_2.paragraph_format.space_before = para.paragraph_format.space_before
            p_2.paragraph_format.space_after = sp_after_body
            p_2.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            para._p.addnext(p_2._p)

            p_3 = doc.add_paragraph(sentence_3)
            p_3.style = para.style
            self._format_gradebook_body_paragraph(p_3, justify=True)
            p_3.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
            p_3.paragraph_format.space_before = para.paragraph_format.space_before
            p_3.paragraph_format.space_after = sp_after_p3
            p_3.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            p_2._p.addnext(p_3._p)
            break


    @staticmethod
    def _compact_gradebook_vertical_legal_textbox(shape):
        """Fit legal text lines inside the narrow vertical textbox."""
        for size in shape.xpath('.//*[local-name()="sz" or local-name()="szCs"]'):
            size.set(qn('w:val'), '10')
        for spacing in shape.xpath('.//*[local-name()="spacing"]'):
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
        for indent in shape.xpath('.//*[local-name()="ind"]'):
            indent.set(qn('w:left'), '0')

    @staticmethod
    def _compact_gradebook_vertical_legal_text(doc):
        """Keep the vertical legal text visible while avoiding wrapping."""
        for shape in doc.element.xpath('.//*[local-name()="txbxContent"]'):
            text = ''.join(
                node.text or '' for node in shape.xpath('.//*[local-name()="t"]')
            )
            if 'Instituto Raimon' not in text or 'B56488687' not in text:
                continue
            IrgCertificateRequest._compact_gradebook_vertical_legal_textbox(shape)

    @staticmethod
    def _restore_gradebook_vertical_legal_text(tpl_path, docx_path):
        """Restore the legal textbox if python-docx drops it on save."""
        with ZipFile(tpl_path) as template_zip:
            template_xml = etree.fromstring(
                template_zip.read('word/document.xml')
            )
        with ZipFile(docx_path) as output_zip:
            output_xml_bytes = output_zip.read('word/document.xml')

        if b'B56488687' in output_xml_bytes and b'B-603323' in output_xml_bytes:
            return

        legal_paragraphs = template_xml.xpath(
            './/*[local-name()="body"]/*[local-name()="p" and '
            './/*[local-name()="t" and contains(text(), "B56488687")]]'
        )
        if not legal_paragraphs:
            return

        legal_runs = [
            child for child in legal_paragraphs[0]
            if etree.QName(child).localname == 'r'
            and child.xpath('.//*[local-name()="txbxContent"]')
        ]
        if not legal_runs:
            return

        output_xml = etree.fromstring(output_xml_bytes)
        first_paragraphs = output_xml.xpath('.//*[local-name()="body"]/*[local-name()="p"]')
        if not first_paragraphs:
            return

        first_paragraph = first_paragraphs[0]
        insert_index = 1 if (
            len(first_paragraph) and etree.QName(first_paragraph[0]).localname == 'pPr'
        ) else 0
        for legal_run in legal_runs:
            for shape in legal_run.xpath('.//*[local-name()="txbxContent"]'):
                IrgCertificateRequest._compact_gradebook_vertical_legal_textbox(shape)
            first_paragraph.insert(insert_index, deepcopy(legal_run))
            insert_index += 1

        tmp_zip = tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False, prefix='cert_gradebook_legal_'
        )
        tmp_zip.close()
        try:
            with ZipFile(docx_path) as source_zip, ZipFile(
                tmp_zip.name, 'w', ZIP_DEFLATED
            ) as target_zip:
                for item in source_zip.infolist():
                    data = source_zip.read(item.filename)
                    if item.filename == 'word/document.xml':
                        data = etree.tostring(
                            output_xml,
                            xml_declaration=True,
                            encoding='UTF-8',
                            standalone=True,
                        )
                    target_zip.writestr(item, data)
            os.replace(tmp_zip.name, docx_path)
        finally:
            if os.path.exists(tmp_zip.name):
                os.unlink(tmp_zip.name)

    @staticmethod
    def _next_relationship_id(rels_xml):
        used_numbers = []
        for rel in rels_xml:
            rel_id = rel.get('Id', '')
            if rel_id.startswith('rId') and rel_id[3:].isdigit():
                used_numbers.append(int(rel_id[3:]))
        return 'rId%s' % ((max(used_numbers) if used_numbers else 0) + 1)

    @staticmethod
    def _next_docpr_id(document_xml):
        ids = []
        for docpr in document_xml.xpath('.//*[local-name()="docPr"]'):
            raw_id = docpr.get('id')
            if raw_id and raw_id.isdigit():
                ids.append(int(raw_id))
        return (max(ids) if ids else 0) + 1

    def _ensure_signature_logo(self, docx_path):
        """Add the logodesgastado.png signature logo next to Raimon Gaja's signature."""
        if self.signer != 'raimon':
            return
        module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        source_image = os.path.join(module_path, 'static', 'src', 'img', 'logodesgastado.png')
        if not os.path.isfile(source_image):
            return

        doc = DocxDocument(docx_path)
        sig_rel_id = None
        for rel_id, rel in doc.part.rels.items():
            target = getattr(rel, 'target_ref', '').lower()
            if any(img in target for img in ('media/image2.jpg', 'media/image2.png', 'media/image2.jpeg')):
                sig_rel_id = rel_id
                break

        if not sig_rel_id:
            return

        paragraphs_to_search = list(doc.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    paragraphs_to_search.extend(cell.paragraphs)

        sig_para = None
        for para in paragraphs_to_search:
            for r in para.runs:
                embed_nodes = r._r.xpath('.//*[@*[local-name()="embed" and .="%s"]]' % sig_rel_id)
                if embed_nodes:
                    sig_para = para
                    break
            if sig_para:
                break

        if sig_para is not None:
            drawing_runs = [r for r in sig_para.runs if r._r.xpath('.//*[local-name()="drawing" or local-name()="pict"]')]
            if len(drawing_runs) < 2:
                sig_para.add_run('       ')
                run_logo = sig_para.add_run()
                run_logo.add_picture(source_image, width=Pt(100))
                doc.save(docx_path)

    def _remove_header_logo(self, docx_path):
        """Remove the header logo/image in word/header1.xml."""
        header_name = 'word/header1.xml'
        with ZipFile(docx_path) as source_zip:
            if header_name not in source_zip.namelist():
                return
            header_xml = etree.fromstring(source_zip.read(header_name))

        runs_to_remove = []
        for r in header_xml.xpath('.//*[local-name()="r"]'):
            if r.xpath('.//*[local-name()="drawing" or local-name()="AlternateContent" or local-name()="pict"]'):
                runs_to_remove.append(r)

        if not runs_to_remove:
            return

        for r in runs_to_remove:
            parent = r.getparent()
            if parent is not None:
                parent.remove(r)

        tmp_zip = tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False, prefix='cert_header_logo_'
        )
        tmp_zip.close()
        try:
            with ZipFile(docx_path) as source_zip, ZipFile(
                tmp_zip.name, 'w', ZIP_DEFLATED
            ) as target_zip:
                for item in source_zip.infolist():
                    data = source_zip.read(item.filename)
                    if item.filename == header_name:
                        data = etree.tostring(
                            header_xml,
                            xml_declaration=True,
                            encoding='UTF-8',
                            standalone=True,
                        )
                    target_zip.writestr(item, data)
            os.replace(tmp_zip.name, docx_path)
        except Exception as e:
            _logger.error("Failed to remove header logo: %s", e)
        finally:
            if os.path.exists(tmp_zip.name):
                os.unlink(tmp_zip.name)

    def _ensure_bottom_right_arcs(self, docx_path):
        """Add the global blue arcs decoration at the page bottom-right."""
        module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        source_image = os.path.join(module_path, 'static', 'src', 'img', 'RCOS.png')
        if not os.path.isfile(source_image):
            return

        rel_ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
        rel_embed = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
        media_name = 'word/media/bottom_right_arcs.png'
        rel_target = 'media/bottom_right_arcs.png'

        with open(source_image, 'rb') as image_file:
            image_data = image_file.read()
        with ZipFile(docx_path) as output_zip:
            output_xml = etree.fromstring(output_zip.read('word/document.xml'))
            output_rels = etree.fromstring(
                output_zip.read('word/_rels/document.xml.rels')
            )

        if output_xml.xpath('.//*[local-name()="docPr" and @name="Bottom Right Arcs"]'):
            return

        output_rel = next(
            (rel for rel in output_rels if rel.get('Target') == rel_target),
            None,
        )
        if output_rel is None:
            rel_id = self._next_relationship_id(output_rels)
            output_rel = etree.Element('{%s}Relationship' % rel_ns)
            output_rel.set('Id', rel_id)
            output_rel.set(
                'Type',
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image',
            )
            output_rel.set('Target', rel_target)
            output_rels.append(output_rel)
        else:
            rel_id = output_rel.get('Id')

        docpr_id = self._next_docpr_id(output_xml)
        drawing_run = etree.fromstring(('''
<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
     xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
     xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
     xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
               relativeHeight="0" behindDoc="1" locked="0" layoutInCell="1"
               allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page"><wp:align>right</wp:align></wp:positionH>
      <wp:positionV relativeFrom="page"><wp:align>bottom</wp:align></wp:positionV>
      <wp:extent cx="3200000" cy="2820000"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="%s" name="Bottom Right Arcs"/>
      <wp:cNvGraphicFramePr/>
      <a:graphic>
        <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <pic:pic>
            <pic:nvPicPr>
              <pic:cNvPr id="%s" name="Bottom Right Arcs"/>
              <pic:cNvPicPr/>
              <pic:nvPr/>
            </pic:nvPicPr>
            <pic:blipFill rotWithShape="1">
              <a:blip r:embed="%s"/>
              <a:stretch><a:fillRect/></a:stretch>
            </pic:blipFill>
            <pic:spPr bwMode="auto">
              <a:xfrm><a:off x="0" y="0"/><a:ext cx="3200000" cy="2820000"/></a:xfrm>
              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
            </pic:spPr>
          </pic:pic>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
</w:r>
''' % (docpr_id, docpr_id, rel_id)).encode('utf-8'))

        body_paragraphs = output_xml.xpath(
            './/*[local-name()="body"]/*[local-name()="p"]'
        )
        if not body_paragraphs:
            return
        body_paragraphs[-1].append(drawing_run)

        tmp_zip = tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False, prefix='cert_arcs_'
        )
        tmp_zip.close()
        try:
            with ZipFile(docx_path) as source_zip, ZipFile(
                tmp_zip.name, 'w', ZIP_DEFLATED
            ) as target_zip:
                existing_names = set(source_zip.namelist())
                for item in source_zip.infolist():
                    data = source_zip.read(item.filename)
                    if item.filename == 'word/document.xml':
                        data = etree.tostring(
                            output_xml,
                            xml_declaration=True,
                            encoding='UTF-8',
                            standalone=True,
                        )
                    elif item.filename == 'word/_rels/document.xml.rels':
                        data = etree.tostring(
                            output_rels,
                            xml_declaration=True,
                            encoding='UTF-8',
                            standalone=True,
                        )
                    target_zip.writestr(item, data)
                if media_name not in existing_names:
                    target_zip.writestr(media_name, image_data)
            os.replace(tmp_zip.name, docx_path)
        finally:
            if os.path.exists(tmp_zip.name):
                os.unlink(tmp_zip.name)

    def _get_certificate_subjects(self):
        """Subjects listed on grade certificates. Extension hook so other
        modules can filter out subjects (e.g. NLEX exemptions)."""
        self.ensure_one()
        return self.gradebook_student_id.gradebook_subject_ids.filtered(
            lambda s: s.op_subject_id.subject_type == 'compulsory'
        )

    @staticmethod
    def _certificate_grade_value(note):
        """Normalize a subject grade for certificate tables.

        Grades strictly below 7 are shown and averaged as 0.
        """
        value = float(note or 0.0)
        return 0.0 if value < 7.0 else value

    @classmethod
    def _format_certificate_grade(cls, note):
        """Return the certificate table display value for a subject grade."""
        return '%.2f' % cls._certificate_grade_value(note)

    def _fill_template(self):
        """Open the .docx template, fill placeholders and table, return bytes."""
        self.ensure_one()
        tpl_path = self._get_template_path()
        if not os.path.isfile(tpl_path):
            raise UserError(
                _('No se encuentra la plantilla Word en %s') % tpl_path
            )

        doc = DocxDocument(tpl_path)
        is_physical = self.certificate_type in PHYSICAL_TYPES
        if is_physical:
            for section in doc.sections:
                section.top_margin = section.top_margin + Pt(75.0)
            # Remove template signature/stamp runs
            sig_rel_ids = []
            for rel_id, rel in doc.part.rels.items():
                target = getattr(rel, 'target_ref', '').lower()
                if any(img in target for img in ('media/image2.jpg', 'media/image2.png', 'media/image2.jpeg')):
                    sig_rel_ids.append(rel_id)
            if sig_rel_ids:
                for para in list(doc.paragraphs):
                    for run in list(para.runs):
                        for rel_id in sig_rel_ids:
                            embed_nodes = run._r.xpath('.//*[@*[local-name()="embed" and .="%s"]]' % rel_id)
                            if embed_nodes:
                                para._p.remove(run._r)
                                break
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for para in list(cell.paragraphs):
                                for run in list(para.runs):
                                    for rel_id in sig_rel_ids:
                                        embed_nodes = run._r.xpath('.//*[@*[local-name()="embed" and .="%s"]]' % rel_id)
                                        if embed_nodes:
                                            para._p.remove(run._r)
                                            break
        scale_percent = 92.5 if is_physical else 75
        self._scale_document_fonts(doc, percent=scale_percent)
        self._replace_dpto_academico_intro(doc)

        top_font_size = None
        for para in doc.paragraphs:
            if para.text.strip():
                for r in para.runs:
                    if r.font and r.font.size:
                        top_font_size = r.font.size
                        break
            if top_font_size:
                break

        if self.document_type == 'gradebook':
            self._compact_gradebook_vertical_legal_text(doc)

        # --- Collect data ---------------------------------------------------
        partner = self.partner_id
        student = self.gradebook_student_id.student_id
        if not student and partner:
            student = self.env['op.student'].sudo().search([('partner_id', '=', partner.id)], limit=1)

        id_label = False
        doc_num = False
        if student:
            if student.document_type_id:
                id_label = student.document_type_id.name
            if student.document_number:
                doc_num = student.document_number

        if not id_label:
            identification_type = getattr(partner, 'l10n_latam_identification_type_id', False)
            id_label = (
                identification_type.name
                if identification_type
                else 'DNI/Pasaporte'
            )
        if not doc_num:
            doc_num = partner.vat or ''

        documento = '%s %s' % (id_label, doc_num)

        subjects = self._get_certificate_subjects()
        nota_media = '%.2f' % (self.gradebook_student_id.total_final or 0.0)

        # Fecha corta DD/MM/YYYY y fecha larga "25 de marzo de 2026"
        fecha = (
            self.request_date.strftime('%d/%m/%Y') if self.request_date else ''
        )
        if self.request_date:
            meses = {
                1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
                5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
                9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre',
            }
            dt = self.request_date
            fecha_larga = '%d de %s de %d' % (dt.day, meses[dt.month], dt.year)
        else:
            fecha_larga = ''

        # --- ECTS and duration (depend on whether course is MNC) ------------
        course_name = self.course_id.name or ''
        is_mnc = _MNC_KEYWORD in course_name
        if self.course_id.id == 4:
            ects_str = '120 ECTS (3000 horas)'
            ects_detallado = '120 ECTS, equivalentes a 3000 horas de estudio'
        else:
            ects_str = '90 ECTS (2250 horas)' if is_mnc else '60 ECTS (1500 horas)'
            ects_detallado = (
                '90 ECTS, equivalentes a 2250 horas de estudio'
                if is_mnc
                else '60 ECTS, equivalentes a 1500 horas de estudio'
            )

        # Academic year range from batch start_date
        batch = self.gradebook_student_id.batch_id
        if batch and batch.start_date:
            start_year = batch.start_date.year
        else:
            start_year = (self.request_date or fields.Datetime.now()).year - 1

        if self.course_id.id == 4:
            end_year = start_year + 2
        else:
            end_year = start_year + (2 if is_mnc else 1)
        periodo_str = '%d-%d' % (start_year, end_year)

        # --- Replace simple placeholders ------------------------------------
        # The Word templates use these exact placeholder names (case-sensitive):
        replacements = {
            '<<NombreAlumno>>': partner.name or '',
            '<<DocumentoIdentidad>>': documento,
            '<<nombreCurso>>': course_name,
            '<<añoCurso>>': periodo_str,
            '<<Etcs>>': ects_str,
            '<<fechaLarga>>': fecha_larga,
            '<<fecha>>': fecha,
            # Remove "modalidad presencial" phrase from template text
            'en la modalidad presencial ': '',
            # Legacy names (keep for backward compatibility)
            '<<nombreAlumno>>': partner.name or '',
            '<<documento>>': documento,
            '<<curso>>': course_name,
            '<<ects>>': ects_str,
            '<<duracion>>': periodo_str,
        }
        if is_physical:
            replacements['Raimon Gaja Jaumeandreu'] = 'Raimon Gaja'
        for para in doc.paragraphs:
            for old, new in replacements.items():
                self._replace_in_paragraph(para, old, new)
        # Also check headers
        for section in doc.sections:
            for para in section.header.paragraphs:
                for old, new in replacements.items():
                    self._replace_in_paragraph(para, old, new)
        # Also check table cells (some placeholders may live inside tables)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old, new in replacements.items():
                            self._replace_in_paragraph(para, old, new)

        has_many_subjects = len(subjects) > 15

        if self.document_type == 'gradebook':
            self._replace_gradebook_description_paragraph(
                doc, partner, documento, course_name, periodo_str, ects_detallado, has_many_subjects=has_many_subjects
            )

        # For non-gradebook types (attendance, enrollment), skip table editing
        if self.document_type not in ('gradebook', 'gradebook_partial'):
            if top_font_size:
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for r in para.runs:
                                    if r.font:
                                        r.font.size = top_font_size
            tmp_docx = tempfile.NamedTemporaryFile(
                suffix='.docx', delete=False, prefix='cert_'
            )
            doc.save(tmp_docx.name)
            tmp_docx.close()
            if self.signer == 'raimon':
                self._ensure_signature_logo(tmp_docx.name)
            self._ensure_bottom_right_arcs(tmp_docx.name)
            return tmp_docx.name

        target_row_height = '235' if has_many_subjects else '315'
        table_font_size = Pt(7.5) if has_many_subjects else (Pt(7.5) if is_physical else (top_font_size or Pt(7.5)))


        if has_many_subjects:
            for section in doc.sections:
                section.top_margin = self._GRADEBOOK_COMPACT_TOP_MARGIN
                section.bottom_margin = self._GRADEBOOK_COMPACT_BOTTOM_MARGIN

            if doc.paragraphs:
                # The first body paragraph is the spacer below the header logo.
                doc.paragraphs[0].paragraph_format.space_after = (
                    self._GRADEBOOK_COMPACT_LOGO_SPACER_AFTER
                )

            # Keep the first paragraph even if empty: it is the logo spacer.
            for p in list(doc.paragraphs)[1:]:
                p_xml = p._p
                drawings = p_xml.findall('.//' + qn('w:drawing')) + p_xml.findall('.//' + qn('w:pict'))
                if not p.text.strip() and not drawings:
                    p_xml.getparent().remove(p_xml)

            for p in doc.paragraphs[1:]:
                txt = p.text.strip()
                p.paragraph_format.line_spacing = 1.0
                if 'CERTIFICA' in txt:
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(8)
                elif 'Las calificaciones obtenidas son:' in txt:
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(10)
                elif 'Para que así conste' in txt:
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(4)
                elif 'Departamento Académico' in txt or 'Raimon Gaja' in txt:
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(2)
                elif txt.startswith('Que '):
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(6)
                else:
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(5)






        # --- Fill the grades table (table index 0) --------------------------
        table = doc.tables[0]
        tbl_xml = table._tbl
        all_rows = tbl_xml.findall(qn('w:tr'))
        # Row 0 = header, rows 1‑12 = data slots, row 13 = Nota Media footer
        header_row = all_rows[0]
        data_rows = all_rows[1:-1]   # rows 1..12
        footer_row = all_rows[-1]

        # Fill data rows with subject info
        for idx, row_xml in enumerate(data_rows):
            cells = row_xml.findall(qn('w:tc'))
            if idx < len(subjects):
                # Normalize row height with hRule="atLeast"
                trPr = row_xml.find(qn('w:trPr'))
                if trPr is None:
                    trPr = row_xml.makeelement(qn('w:trPr'), {})
                    row_xml.insert(0, trPr)
                for h in trPr.findall(qn('w:trHeight')):
                    trPr.remove(h)
                trHeight = trPr.makeelement(qn('w:trHeight'), {
                    qn('w:val'): target_row_height,
                    qn('w:hRule'): 'atLeast',
                })
                trPr.append(trHeight)

                subj = subjects[idx]
                cell_values = [
                    subj.op_subject_id.code or '',
                    subj.op_subject_id.name or '',
                    self._format_certificate_grade(subj.final_subject_note),
                ]
                for ci, val in enumerate(cell_values):
                    if ci < len(cells):
                        for p in cells[ci].findall(qn('w:p')):
                            for r in p.findall(qn('w:r')):
                                t = r.find(qn('w:t'))
                                if t is not None:
                                    t.text = val
                                    break
                            else:
                                # No run exists — create one
                                r_el = p.makeelement(qn('w:r'), {})
                                rpr_src = data_rows[0].findall(qn('w:tc'))[0] \
                                    .findall(qn('w:p'))[0].findall(qn('w:r'))
                                if rpr_src:
                                    rpr = rpr_src[0].find(qn('w:rPr'))
                                    if rpr is not None:
                                        r_el.append(deepcopy(rpr))
                                t_el = r_el.makeelement(qn('w:t'), {})
                                t_el.text = val
                                r_el.append(t_el)
                                p.append(r_el)
                            break  # only first <w:p>
            else:
                # More rows than subjects → remove the excess row
                tbl_xml.remove(row_xml)

        # If there are more subjects than template rows (>12), clone rows
        if len(subjects) > len(data_rows):
            ref_row = data_rows[0]  # use first data row as formatting reference
            for idx in range(len(data_rows), len(subjects)):
                subj = subjects[idx]
                new_row = deepcopy(ref_row)
                # Normalize row height with hRule="atLeast"
                trPr = new_row.find(qn('w:trPr'))
                if trPr is None:
                    trPr = new_row.makeelement(qn('w:trPr'), {})
                    new_row.insert(0, trPr)
                for h in trPr.findall(qn('w:trHeight')):
                    trPr.remove(h)
                trHeight = trPr.makeelement(qn('w:trHeight'), {
                    qn('w:val'): target_row_height,
                    qn('w:hRule'): 'atLeast',
                })
                trPr.append(trHeight)
                cells = new_row.findall(qn('w:tc'))
                cell_values = [
                    subj.op_subject_id.code or '',
                    subj.op_subject_id.name or '',
                    self._format_certificate_grade(subj.final_subject_note),
                ]
                for ci, val in enumerate(cell_values):
                    if ci < len(cells):
                        for p in cells[ci].findall(qn('w:p')):
                            for r in p.findall(qn('w:r')):
                                t = r.find(qn('w:t'))
                                if t is not None:
                                    t.text = val
                                    break
                            break
                footer_row.addprevious(new_row)

        # Normalize cell top and bottom borders and remove shading for all data rows
        all_data_rows = tbl_xml.findall(qn('w:tr'))[1:-1]
        for r_idx, r_xml in enumerate(all_data_rows):
            for c in r_xml.findall(qn('w:tc')):
                tcPr = c.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = c.makeelement(qn('w:tcPr'), {})
                    c.insert(0, tcPr)
                for shd in tcPr.findall(qn('w:shd')):
                    tcPr.remove(shd)
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
                    tcPr.append(tcBorders)
                for b_name in ('top', 'bottom'):
                    for b in tcBorders.findall(qn(f'w:{b_name}')):
                        tcBorders.remove(b)
                t_elem = tcBorders.makeelement(qn('w:top'), {
                    qn('w:val'): 'single',
                    qn('w:color'): 'dee2e6',
                    qn('w:sz'): '5',
                    qn('w:space'): '0',
                })
                tcBorders.append(t_elem)
                if r_idx < len(all_data_rows) - 1:
                    b_elem = tcBorders.makeelement(qn('w:bottom'), {
                        qn('w:val'): 'single',
                        qn('w:color'): 'dee2e6',
                        qn('w:sz'): '5',
                        qn('w:space'): '0',
                    })
                else:
                    b_elem = tcBorders.makeelement(qn('w:bottom'), {
                        qn('w:val'): 'single',
                        qn('w:color'): '000000',
                        qn('w:sz'): '10',
                        qn('w:space'): '0',
                    })
                tcBorders.append(b_elem)


        # Fill Nota Media in footer row.
        # The footer row may have merged cells; we write the grade in the last
        # <w:tc> element.  We also attempt to write into every cell that does
        # NOT contain the label "Nota Media" to cover different merge layouts.
        footer_cells = footer_row.findall(qn('w:tc'))
        nota_written = False
        for cell in reversed(footer_cells):
            # Read existing text of this cell to decide if it's the label cell
            cell_text = ''.join(
                t.text or ''
                for p in cell.findall(qn('w:p'))
                for r in p.findall(qn('w:r'))
                for t in r.findall(qn('w:t'))
            ).strip()
            if 'Nota Media' in cell_text or 'nota media' in cell_text.lower():
                continue  # skip the label cell
            # This is the value cell — fill it
            for p in cell.findall(qn('w:p')):
                runs = p.findall(qn('w:r'))
                if runs:
                    for r in runs:
                        t = r.find(qn('w:t'))
                        if t is not None:
                            t.text = nota_media
                            nota_written = True
                            break
                else:
                    r_el = p.makeelement(qn('w:r'), {})
                    t_el = r_el.makeelement(qn('w:t'), {})
                    t_el.text = nota_media
                    r_el.append(t_el)
                    p.append(r_el)
                    nota_written = True
                break
            if nota_written:
                break

        # Fallback: if the footer row has only one cell (fully merged), append
        # the nota media after the label text.
        if not nota_written and footer_cells:
            last_cell = footer_cells[-1]
            paras = last_cell.findall(qn('w:p'))
            target_p = paras[-1] if paras else None
            if target_p is not None:
                r_el = target_p.makeelement(qn('w:r'), {})
                t_el = r_el.makeelement(qn('w:t'), {})
                t_el.text = '  ' + nota_media
                r_el.append(t_el)
                target_p.append(r_el)

        if self.document_type == 'gradebook':
            self._format_gradebook_static_paragraphs(doc)

        if is_physical:
            for para in doc.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9.25)

        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for r in para.runs:
                            if r.font:
                                r.font.size = table_font_size

        # Save filled document to a temp file
        tmp_docx = tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False, prefix='cert_'
        )
        doc.save(tmp_docx.name)
        tmp_docx.close()
        if self.document_type == 'gradebook':
            self._restore_gradebook_vertical_legal_text(tpl_path, tmp_docx.name)
        if self.signer == 'raimon' and not is_physical:
            self._ensure_signature_logo(tmp_docx.name)
        if self.document_type == 'gradebook' and self.certificate_type in PHYSICAL_TYPES:
            self._remove_header_logo(tmp_docx.name)
        else:
            self._ensure_bottom_right_arcs(tmp_docx.name)
        return tmp_docx.name

    @staticmethod
    def _convert_to_pdf(docx_path):
        """Convert a .docx file to PDF using LibreOffice and return PDF bytes."""
        out_dir = os.path.dirname(docx_path)
        try:
            subprocess.run(
                [
                    'libreoffice', '--headless', '--norestore',
                    '--convert-to', 'pdf',
                    '--outdir', out_dir,
                    docx_path,
                ],
                check=True,
                timeout=60,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
        except FileNotFoundError:
            raise UserError(
                _('LibreOffice no está instalado en el servidor. '
                  'Ejecute: apt-get install -y libreoffice-writer')
            )
        except subprocess.CalledProcessError as exc:
            _logger.error('LibreOffice conversion failed: %s', exc.stderr)
            raise UserError(
                _('Error al convertir el certificado a PDF. Revise el log.')
            )

        pdf_path = docx_path.rsplit('.', 1)[0] + '.pdf'
        if not os.path.isfile(pdf_path):
            raise UserError(_('No se generó el archivo PDF.'))

        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()

        # Cleanup temp files
        for path in (docx_path, pdf_path):
            try:
                os.unlink(path)
            except OSError:
                pass

        return pdf_bytes

    def _generate_diploma_pdf_content(self):
        self.ensure_one()
        # Get sequence
        registry_number = self.env['ir.sequence'].next_by_code('irg.diploma.registry') or 'DRAFT'
        
        # Format dates
        from babel.dates import format_date
        from urllib.parse import urlencode
        
        date_to_use = self.request_date or fields.Date.context_today(self)
        # Convert datetime to date for formatting
        date_val = date_to_use.date() if hasattr(date_to_use, 'date') else date_to_use
        date_es = "{} de {} de {}".format(
            date_val.day, 
            format_date(date_val, format='MMMM', locale='es_ES'), 
            date_val.year
        )
        date_cat = "{} de {} de {}".format(
            date_val.day, 
            format_date(date_val, format='MMMM', locale='ca_ES'), 
            date_val.year
        )

        # Get names
        student_name = self.partner_id.name or ""
        course_name_es = self.course_id.name or ""
        course_name_cat = getattr(self.course_id, 'name_cat', None) or course_name_es
        
        pdf_generator = self.env['report.irg_generacion_diplomas.diploma_pdf']
        course_name_cat = pdf_generator._normalize_catalan_course_name(course_name_cat)

        # QR URL
        query_params = {'id': registry_number}
        student = self.gradebook_student_id.student_id
        if not student:
            student = self.env['op.student'].search([('partner_id', '=', self.partner_id.id)], limit=1)
            
        if student and 'op.sign_certificate' in self.env:
            stamp_payload = {
                'registry_number': registry_number,
                'student_name': student_name,
                'course_name_es': course_name_es,
                'course_name_cat': course_name_cat,
                'issue_date': str(date_val),
                'diploma_type': 'digital' if self.certificate_type in ('digital', 'custom') else 'physical',
            }
            stamp_data = self.env['op.sign_certificate'].sudo().stamp_data(stamp_payload, student=student) or {}
            if stamp_data.get('stamp') and stamp_data.get('data_str') and stamp_data.get('certificate_id'):
                query_params.update({
                    'stamp': stamp_data.get('stamp'),
                    'data_str': stamp_data.get('data_str'),
                    'certificate_id': stamp_data.get('certificate_id'),
                })

        qr_url = "https://institutoraimongaja.com/verificar/?{}".format(urlencode(query_params))

        def html_split(name, lang='es'):
            if not name:
                return name
            sep = ' y ' if lang == 'es' else ' i '
            if sep in name:
                parts = name.rsplit(sep, 1)
                return parts[0] + sep.strip() + '<br/>' + parts[1]
            return name

        data = {
            'student_name': student_name,
            'course_name_es': course_name_es,
            'course_name_cat': course_name_cat,
            'course_name_es_html': html_split(course_name_es, lang='es'),
            'course_name_cat_html': html_split(course_name_cat, lang='cat'),
            'date_es': date_es,
            'date_cat': date_cat,
            'registry_number': registry_number,
            'qr_url': qr_url,
        }
        
        diploma_type = 'physical' if self.certificate_type in PHYSICAL_TYPES else 'digital'
        pdf_content = pdf_generator.generate_diploma_pdf(data, diploma_type=diploma_type)
        
        # Create irg.diploma.registry record
        student_course = self.env['op.student.course'].search([
            ('student_id', '=', student.id if student else False),
            ('course_id', '=', self.course_id.id),
        ], limit=1)
        
        self.env['irg.diploma.registry'].sudo().create({
            'registry_number': registry_number,
            'student_id': student.id if student else False,
            'student_course_id': student_course.id if student_course else False,
            'issue_date': date_val,
            'diploma_type': diploma_type,
            'qr_url': qr_url,
            'state': 'valid',
        })
        
        return pdf_content

    def _generate_and_attach_pdf(self):
        """Fill the Word template or generate ReportLab diploma, and convert/attach as PDF."""
        self.ensure_one()
        try:
            if self.document_type == 'diploma':
                pdf_content = self._generate_diploma_pdf_content()
            else:
                docx_path = self._fill_template()
                pdf_content = self._convert_to_pdf(docx_path)
        except UserError:
            raise
        except Exception as exc:
            _logger.error(
                'Error generando PDF para certificado %s: %s', self.name, exc,
                exc_info=True,
            )
            raise UserError(
                _('No se pudo generar el PDF del certificado. Revise el log.')
            )

        filename = 'Certificado_%s_%s.pdf' % (
            (self.partner_id.name or 'Alumno').replace(' ', '_'),
            self.name,
        )
        attachment = self.env['ir.attachment'].sudo().create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(pdf_content),
            'res_model': self._name,
            'res_id': self.id,
            'mimetype': 'application/pdf',
            'public': False,
        })
        self.attachment_id = attachment

    def action_generate_pdf(self):
        """Backend button to (re)generate the PDF."""
        self._generate_and_attach_pdf()

    def action_download_pdf(self):
        """Return file-download action for the attached PDF."""
        self.ensure_one()
        if not self.attachment_id:
            raise UserError(_('No hay PDF generado aún. Genera el certificado primero.'))
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/content/%d?download=true' % self.attachment_id.id,
            'target': 'self',
        }

    # ------------------------------------------------------------------
    # Email notifications
    # ------------------------------------------------------------------

    def _send_digital_notification(self):
        template = self.env.ref(
            'irg_gradebook_certificates.mail_template_cert_digital_ready',
            raise_if_not_found=False,
        )
        if template:
            # No force_send to avoid blocking; use queue
            template.send_mail(self.id, force_send=False)

    def _send_paid_notification(self):
        template = self.env.ref(
            'irg_gradebook_certificates.mail_template_cert_paid_physical',
            raise_if_not_found=False,
        )
        if template:
            template.send_mail(self.id, force_send=False)

    def _send_sent_notification(self):
        template = self.env.ref(
            'irg_gradebook_certificates.mail_template_cert_sent',
            raise_if_not_found=False,
        )
        if template:
            template.send_mail(self.id, force_send=False)

    def _send_team_notification(self):
        """Notify the academic team that a physical certificate has been paid."""
        template = self.env.ref(
            'irg_gradebook_certificates.mail_template_cert_physical_team',
            raise_if_not_found=False,
        )
        if template:
            template.send_mail(self.id, force_send=False)

    # ------------------------------------------------------------------
    # Portal invoice + payment link
    # ------------------------------------------------------------------

    def _create_portal_invoice(self):
        """Create an out_invoice for this certificate request and store a
        payment link on the record.

        Called by the portal controller immediately after the certificate
        request is created.  Stores the posted invoice in self.invoice_id
        and the payment URL in self.payment_link.
        """
        self.ensure_one()
        partner = self.partner_id

        # Build invoice lines
        cert_tmpl = self.env.ref(
            _PORTAL_PRODUCT_XMLID[self.certificate_type]
        ).sudo()
        cert_product = cert_tmpl.product_variant_ids[:1]
        if not cert_product:
            raise UserError(
                _('El producto del certificado no está configurado correctamente.')
            )

        invoice_lines = [(0, 0, {
            'product_id': cert_product.id,
            'quantity': 1,
            'price_unit': PRICE_MAP.get(self.certificate_type, 0.0),
            'name': cert_tmpl.name,
        })]

        if self.certificate_type in PHYSICAL_TYPES and self.shipping_type:
            ship_tmpl = self.env.ref(
                _PORTAL_SHIPPING_XMLID[self.shipping_type]
            ).sudo()
            ship_product = ship_tmpl.product_variant_ids[:1]
            if ship_product:
                invoice_lines.append((0, 0, {
                    'product_id': ship_product.id,
                    'quantity': 1,
                    'price_unit': SHIPPING_MAP.get(self.shipping_type, 0.0),
                    'name': ship_tmpl.name,
                }))

        # sudo() needed: called from portal context where the user lacks
        # account.move creation rights.
        invoice = self.env['account.move'].sudo().create({
            'move_type': 'out_invoice',
            'partner_id': partner.id,
            'invoice_line_ids': invoice_lines,
            'narration': _('Certificado de notas %s \u2014 %s') % (
                dict(CERTIFICATE_TYPES).get(self.certificate_type, ''),
                self.name,
            ),
            'ref': self.name,
        })
        invoice.sudo().action_post()
        self.invoice_id = invoice

        # Generate payment link via payment.link.wizard.
        # sudo() needed: payment models are restricted to internal users.
        try:
            wizard = self.env['payment.link.wizard'].with_context(
                active_model='account.move',
                active_id=invoice.id,
            ).sudo().create({
                'res_model': 'account.move',
                'res_id': invoice.id,
                'amount': invoice.amount_residual,
                'amount_max': invoice.amount_residual,
                'currency_id': invoice.currency_id.id,
                'partner_id': partner.id,
                'description': _('Certificado %s') % self.name,
            })
            self.payment_link = wizard.link
        except Exception:
            _logger.warning(
                'No se pudo generar el enlace de pago para el certificado %s; '
                'el alumno podrá pagar desde el portal de facturas.',
                self.name, exc_info=True,
            )
            # Fallback: portal URL of the invoice with access token
            invoice._portal_ensure_token()
            self.payment_link = invoice.get_portal_url()

    # ------------------------------------------------------------------
    # Cron: auto-process paid portal certificates
    # ------------------------------------------------------------------

    @api.model
    def _cron_process_paid_certificates(self):
        """Detect portal certificate requests whose invoice has been paid
        and trigger _process_payment() on each one.

        Called every 10 minutes by the ir.cron record defined in
        data/cron_data.xml.
        """
        pending = self.search([
            ('state', '=', 'pending_payment'),
            ('origin', '=', 'portal'),
            ('invoice_id', '!=', False),
        ])
        for cert in pending:
            if cert.invoice_id.payment_state in ('paid', 'in_payment'):
                try:
                    cert.sudo()._process_payment()
                except Exception:
                    _logger.exception(
                        'Error procesando certificado %s tras pago de factura %s',
                        cert.name,
                        cert.invoice_id.name,
                    )
