# -*- coding: utf-8 -*-
import os
import tempfile
import logging
from zipfile import ZipFile, ZIP_DEFLATED
from copy import deepcopy
from docx import Document as DocxDocument
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree

from odoo import models, fields, api, _
from odoo.exceptions import UserError
from odoo.modules.module import get_module_resource

_logger = logging.getLogger(__name__)

class IrgCertificateRequest(models.Model):
    _inherit = 'irg.certificate.request'

    _PARTIAL_TEXT_INDENT = Twips(-172)
    _PARTIAL_TABLE_WIDTH = Twips(9010)
    _PARTIAL_PAGE_TEXT_WIDTH = Twips(8055)
    _PARTIAL_TEXT_RIGHT_INDENT = (
        _PARTIAL_PAGE_TEXT_WIDTH - _PARTIAL_TEXT_INDENT - _PARTIAL_TABLE_WIDTH
    )
    _DPTO_ACADEMICO_INTRO = (
        'El Instituto Raimon Gaja, con CIF B-56488687 en calle '
        'Córcega 213, 1º 2ª, 08036 Barcelona.'
    )

    @staticmethod
    def _replace_paragraph_text_with_bold_segments(paragraph, segments):
        """Replace paragraph text with runs, preserving first-run style.

        ``segments`` is an iterable of ``(text, bold)`` pairs. Empty text
        fragments are ignored so the resulting paragraph only contains the
        content that must be rendered.
        """
        base_run = paragraph.runs[0] if paragraph.runs else None
        base_rpr = None
        if base_run is not None and base_run._r.rPr is not None:
            base_rpr = deepcopy(base_run._r.rPr)

        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)

        for text, bold in segments:
            if not text:
                continue
            run = paragraph.add_run(text)
            if base_rpr is not None:
                run._r.insert(0, deepcopy(base_rpr))
            run.bold = bool(bold)

    def _format_partial_body_paragraph(self, paragraph, justify=True):
        """Align body text with the grade table width.

        The official notes table has a fixed width and a small negative table
        indent. Matching the surrounding paragraphs to that geometry keeps the
        text visually constrained to the same amplitude as the table.
        """
        fmt = paragraph.paragraph_format
        fmt.left_indent = self._PARTIAL_TEXT_INDENT
        fmt.right_indent = self._PARTIAL_TEXT_RIGHT_INDENT
        if justify:
            paragraph.alignment = 3  # WD_ALIGN_PARAGRAPH.JUSTIFY
        return paragraph

    def _format_partial_signature_paragraph(self, paragraph):
        """Normalize signature lines to the same left-aligned text grid."""
        is_physical = self.certificate_type in ('physical', 'physical_apostilled')
        normalized_text = ' '.join(paragraph.text.split())
        if normalized_text == 'Departamento Académico Instituto Raimon Gaja':
            paragraph.text = 'Departamento Académico\nInstituto Raimon Gaja'
        elif normalized_text in (
            'Raimon Gaja Jaumeandreu Instituto Raimon Gaja',
            'Raimon Gaja Instituto Raimon Gaja'
        ):
            if is_physical:
                paragraph.text = 'Raimon Gaja\nDirector General iRG'
            else:
                paragraph.text = 'Raimon Gaja Jaumeandreu\nInstituto Raimon Gaja'
        self._format_partial_body_paragraph(paragraph, justify=False)
        paragraph.alignment = 0  # WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.tab_stops.clear_all()
        for run in paragraph.runs:
            if run.text:
                run.text = run.text.lstrip()
            if is_physical:
                run.font.size = Pt(8.5)
        return paragraph

    def _replace_dpto_academico_intro(self, doc):
        """Apply the requested issuer sentence for academic department signer."""
        if self.signer != 'dpto_academico':
            return
        for para in doc.paragraphs:
            text = ''.join(run.text for run in para.runs).strip()
            if text == 'El Departamento Académico del Instituto Raimon Gaja, S.L.':
                para.text = self._DPTO_ACADEMICO_INTRO
                self._format_partial_body_paragraph(para, justify=True)
                break

    def _format_partial_closing_paragraphs(self, doc):
        """Justify closing text and constrain it to the grade table width."""
        is_physical = self.certificate_type in ('physical', 'physical_apostilled')
        for para in doc.paragraphs:
            if 'Para que así conste' in para.text:
                if is_physical:
                    fecha_larga = ''
                    if self.request_date:
                        meses = {
                            1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
                            5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
                            9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre',
                        }
                        dt = self.request_date
                        fecha_larga = '%d de %s de %d' % (dt.day, meses[dt.month], dt.year)
                    para.text = f'Para que así conste, firmo la presente en Barcelona, a fecha {fecha_larga}'
                    para.paragraph_format.space_after = Pt(48)
                    for run in para.runs:
                        run.font.size = Pt(8.5)
                self._format_partial_body_paragraph(para, justify=True)

    def _format_partial_static_paragraphs(self, doc):
        """Align fixed template paragraphs with the notes table grid."""
        is_physical = self.certificate_type in ('physical', 'physical_apostilled')
        signer_intro_markers = (
            'Raimon Gaja Jaumeandreu, con DNI',
            'Raimon Gaja, con DNI',
            self._DPTO_ACADEMICO_INTRO,
        )
        signature_texts = (
            'Raimon Gaja Jaumeandreu Instituto Raimon Gaja',
            'Raimon Gaja Instituto Raimon Gaja',
            'Departamento Académico Instituto Raimon Gaja',
        )
        for para in doc.paragraphs:
            text = para.text.strip()
            normalized_text = ' '.join(text.split())
            if not text:
                continue
            if (
                text == 'CERTIFICA:'
                or any(marker in text for marker in signer_intro_markers)
            ):
                self._format_partial_body_paragraph(para, justify=False)
                para.alignment = 0  # WD_ALIGN_PARAGRAPH.LEFT
            if normalized_text in signature_texts:
                self._format_partial_signature_paragraph(para)

    @staticmethod
    def _compact_vertical_legal_text(doc):
        """Keep the vertical legal text visible while avoiding margin clipping."""
        for shape in doc.element.xpath('.//*[local-name()="txbxContent"]'):
            text = ''.join(
                node.text or '' for node in shape.xpath('.//*[local-name()="t"]')
            )
            if 'Instituto Raimon' not in text or 'B56488687' not in text:
                continue
            IrgCertificateRequest._compact_vertical_legal_textbox(shape)

    @staticmethod
    def _compact_vertical_legal_textbox(shape):
        """Fit legal text lines inside the narrow vertical textbox."""
        for size in shape.xpath('.//*[local-name()="sz" or local-name()="szCs"]'):
            size.set(qn('w:val'), '10')
        for spacing in shape.xpath('.//*[local-name()="spacing"]'):
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
        for indent in shape.xpath('.//*[local-name()="ind"]'):
            indent.set(qn('w:left'), '0')

    @staticmethod
    def _restore_vertical_legal_text(tpl_path, docx_path):
        """Restore the legal text box if python-docx drops it on save."""
        with ZipFile(tpl_path) as template_zip:
            template_xml = etree.fromstring(
                template_zip.read('word/document.xml')
            )
        with ZipFile(docx_path) as output_zip:
            output_xml_bytes = output_zip.read('word/document.xml')

        if b'B56488687' in output_xml_bytes and b'B-603323' in output_xml_bytes:
            return

        legal_runs = []
        legal_paragraphs = template_xml.xpath(
            './/*[local-name()="body"]/*[local-name()="p" and '
            './/*[local-name()="t" and contains(text(), "B56488687")]]'
        )
        if legal_paragraphs:
            legal_runs = [
                child for child in legal_paragraphs[0]
                if etree.QName(child).localname == 'r'
                and child.xpath('.//*[local-name()="txbxContent"]')
            ]

        if not legal_runs:
            return

        output_xml = etree.fromstring(output_xml_bytes)
        first_paragraphs = output_xml.xpath('.//*[local-name()="body"]/*[local-name()="p"]')
        if not first_paragraphs:
            return

        first_paragraph = first_paragraphs[0]
        insert_index = 1 if (
            len(first_paragraph) and etree.QName(first_paragraph[0]).localname == 'pPr'
        ) else 0
        for legal_run in legal_runs:
            for shape in legal_run.xpath('.//*[local-name()="txbxContent"]'):
                IrgCertificateRequest._compact_vertical_legal_textbox(shape)
            first_paragraph.insert(insert_index, deepcopy(legal_run))
            insert_index += 1

        tmp_zip = tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False, prefix='cert_partial_legal_'
        )
        tmp_zip.close()
        try:
            with ZipFile(docx_path) as source_zip, ZipFile(
                tmp_zip.name, 'w', ZIP_DEFLATED
            ) as target_zip:
                for item in source_zip.infolist():
                    data = source_zip.read(item.filename)
                    if item.filename == 'word/document.xml':
                        data = etree.tostring(
                            output_xml,
                            xml_declaration=True,
                            encoding='UTF-8',
                            standalone=True,
                        )
                    target_zip.writestr(item, data)
            os.replace(tmp_zip.name, docx_path)
        finally:
            if os.path.exists(tmp_zip.name):
                os.unlink(tmp_zip.name)



    def _get_template_path(self):
        if self.document_type == 'gradebook_partial':
            signer_suffix = 'dpto' if self.signer == 'dpto_academico' else 'raimon'
            filename = f'Plantilla-certificado-notas-{signer_suffix}.docx'
            path = get_module_resource('irg_gradebook_certificates', 'static', 'src', 'templates', filename)
            if path:
                return path
        return super()._get_template_path()

    def _fill_template(self):
        if self.document_type != 'gradebook_partial':
            return super()._fill_template()

        self.ensure_one()
        tpl_path = self._get_template_path()
        if not tpl_path or not os.path.isfile(tpl_path):
            raise UserError(
                _('No se encuentra la plantilla Word en %s') % tpl_path
            )

        doc = DocxDocument(tpl_path)
        is_physical = self.certificate_type in ('physical', 'physical_apostilled')
        if is_physical:
            for section in doc.sections:
                section.top_margin = section.top_margin + Pt(37.5)
            # Remove template signature/stamp runs
            sig_rel_ids = []
            for rel_id, rel in doc.part.rels.items():
                target = getattr(rel, 'target_ref', '').lower()
                if any(img in target for img in ('media/image2.jpg', 'media/image2.png', 'media/image2.jpeg')):
                    sig_rel_ids.append(rel_id)
            if sig_rel_ids:
                for para in list(doc.paragraphs):
                    for run in list(para.runs):
                        for rel_id in sig_rel_ids:
                            embed_nodes = run._r.xpath('.//*[@*[local-name()="embed" and .="%s"]]' % rel_id)
                            if embed_nodes:
                                para._p.remove(run._r)
                                break
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for para in list(cell.paragraphs):
                                for run in list(para.runs):
                                    for rel_id in sig_rel_ids:
                                        embed_nodes = run._r.xpath('.//*[@*[local-name()="embed" and .="%s"]]' % rel_id)
                                        if embed_nodes:
                                            para._p.remove(run._r)
                                            break
        scale_percent = 85 if is_physical else 75
        self._scale_document_fonts(doc, percent=scale_percent)

        top_font_size = None
        for para in doc.paragraphs:
            if para.text.strip():
                for r in para.runs:
                    if r.font and r.font.size:
                        top_font_size = r.font.size
                        break
            if top_font_size:
                break

        self._replace_dpto_academico_intro(doc)
        self._compact_vertical_legal_text(doc)

        # --- Collect data ---------------------------------------------------
        partner = self.partner_id
        identification_type = getattr(partner, 'l10n_latam_identification_type_id', False)
        id_label = (
            identification_type.name
            if identification_type
            else 'DNI/Pasaporte'
        )
        documento = '%s %s' % (id_label, partner.vat or '')

        # Filtrar asignaturas obligatorias (hook extensible, p.ej. exclusion NLEX)
        subjects = self._get_certificate_subjects()

        subject_notes = []
        valid_notes = []
        for subj in subjects:
            gradebook_info = subj._get_gradebook_info(subj)
            qty_configured = gradebook_info.get('exam', {}).get('qty', 0) or 0
            exam_results = subj.gradebook_result_ids.filtered(lambda r: r.survey_type == 'exam')
            
            # Si una asignatura no tiene exámenes calificados (exam_results vacío) 
            # o la cantidad de exámenes es inferior al qty configurado, se escribe Pendiente
            if not exam_results or len(exam_results) < qty_configured:
                note_str = 'Pendiente'
            else:
                note_str = '%.2f' % (subj.final_subject_note or 0.0)
                valid_notes.append(subj.final_subject_note or 0.0)
                
            subject_notes.append({
                'code': subj.op_subject_id.code or '',
                'name': subj.op_subject_id.name or '',
                'note': note_str,
            })

        # Calcular la nota media del certificado usando solo las notas no pendientes
        if valid_notes:
            nota_media = '%.2f' % (sum(valid_notes) / len(valid_notes))
        else:
            nota_media = 'Pendiente'

        # Fecha corta DD/MM/YYYY y fecha larga
        fecha = (
            self.request_date.strftime('%d/%m/%Y') if self.request_date else ''
        )
        if self.request_date:
            meses = {
                1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
                5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
                9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre',
            }
            dt = self.request_date
            fecha_larga = '%d de %s de %d' % (dt.day, meses[dt.month], dt.year)
        else:
            fecha_larga = ''

        course_name = self.course_id.name or ''
        is_mnc = 'MNC' in course_name
        if self.course_id.id == 4:
            ects_str = '120 ECTS (3000 horas)'
            ects_detallado = '120 ECTS, equivalentes a 3000 horas de estudio'
        else:
            ects_str = '90 ECTS (2250 horas)' if is_mnc else '60 ECTS (1500 horas)'
            ects_detallado = '90 ECTS, equivalentes a 2250 horas de estudio' if is_mnc else '60 ECTS, equivalentes a 1500 horas de estudio'

        batch = self.gradebook_student_id.batch_id
        if batch and batch.start_date:
            start_year = batch.start_date.year
        else:
            start_year = (self.request_date or fields.Datetime.now()).year - 1

        if self.course_id.id == 4:
            end_year = start_year + 2
        else:
            end_year = start_year + (2 if is_mnc else 1)
        periodo_str = '%d-%d' % (start_year, end_year)

        # Determinar género del estudiante para "matriculado/a"
        gender_word = 'consta matriculado/a'
        student = self.gradebook_student_id.student_id
        if student and student.gender:
            if student.gender == 'f':
                gender_word = 'consta matriculada'
            elif student.gender == 'm':
                gender_word = 'consta matriculado'

        # Match the final gradebook certificate: print the exact partner ID type.
        documento_formateado = documento

        sentence_1 = 'Que %s con %s %s en el %s durante el período académico %s.' % (
            partner.name or '',
            documento_formateado,
            gender_word,
            course_name,
            periodo_str
        )
        sentence_2 = 'Que, el Máster consta de %s, distribuidas entre horas de clases y horas destinadas a otras actividades académicas.' % (
            ects_detallado
        )
        sentence_3 = 'Las calificaciones obtenidas son:'

        target_text = 'Que <<NombreAlumno>> con <<DocumentoIdentidad>> matriculado/a en el <<nombreCurso>> impartido en la modalidad presencial durante el periodo académico <<añoCurso>> con una carga lectiva de <<Etcs>>, ha obtenido las calificaciones siguientes:'

        # Modificar "CERTIFICA" a "CERTIFICA:" para coincidir con el diseño solicitado
        for para in list(doc.paragraphs):
            full_text = ''.join(r.text for r in para.runs).strip()
            if full_text == 'CERTIFICA':
                para.text = 'CERTIFICA:'
                self._format_partial_body_paragraph(para, justify=False)
                para.alignment = 0  # WD_ALIGN_PARAGRAPH.LEFT

        for para in list(doc.paragraphs):
            full_text = ''.join(r.text for r in para.runs)
            if target_text in full_text:
                # Reemplazar el primer párrafo con la primera frase, alineación a la izquierda y espaciado
                self._replace_paragraph_text_with_bold_segments(para, [
                    ('Que ', False),
                    (partner.name or '', True),
                    (' con %s %s en el ' % (documento_formateado, gender_word), False),
                    (course_name, True),
                    (' durante el período académico %s.' % periodo_str, False),
                ])
                self._format_partial_body_paragraph(para, justify=True)
                para.paragraph_format.space_after = Pt(12)
                
                # Crear el segundo párrafo copiando estilo, márgenes y alineación justificada
                p_2 = doc.add_paragraph(sentence_2)
                p_2.style = para.style
                self._format_partial_body_paragraph(p_2, justify=True)
                p_2.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                p_2.paragraph_format.space_before = para.paragraph_format.space_before
                p_2.paragraph_format.space_after = Pt(12)
                p_2.paragraph_format.line_spacing = para.paragraph_format.line_spacing
                para._p.addnext(p_2._p)
                
                # Crear el tercer párrafo copiando estilo, márgenes y alineación justificada
                p_3 = doc.add_paragraph(sentence_3)
                p_3.style = para.style
                self._format_partial_body_paragraph(p_3, justify=True)
                p_3.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                p_3.paragraph_format.space_before = para.paragraph_format.space_before
                p_3.paragraph_format.space_after = Pt(12)
                p_3.paragraph_format.line_spacing = para.paragraph_format.line_spacing
                p_2._p.addnext(p_3._p)
                break

        replacements = {
            '<<NombreAlumno>>': partner.name or '',
            '<<DocumentoIdentidad>>': documento,
            '<<nombreCurso>>': course_name,
            '<<añoCurso>>': periodo_str,
            '<<Etcs>>': ects_str,
            '<<fechaLarga>>': fecha_larga,
            '<<fecha>>': fecha,
            '<<nombreAlumno>>': partner.name or '',
            '<<documento>>': documento,
            '<<curso>>': course_name,
            '<<ects>>': ects_str,
            '<<duracion>>': periodo_str,
        }
        if is_physical:
            replacements['Raimon Gaja Jaumeandreu'] = 'Raimon Gaja'

        for para in doc.paragraphs:
            for old, new in replacements.items():
                self._replace_in_paragraph(para, old, new)
        self._format_partial_closing_paragraphs(doc)
        self._format_partial_static_paragraphs(doc)
        for section in doc.sections:
            for para in section.header.paragraphs:
                for old, new in replacements.items():
                    self._replace_in_paragraph(para, old, new)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old, new in replacements.items():
                            self._replace_in_paragraph(para, old, new)

        # Rellenar la tabla de notas
        table = doc.tables[0]
        tbl_xml = table._tbl
        all_rows = tbl_xml.findall(qn('w:tr'))
        data_rows = all_rows[1:-1]
        footer_row = all_rows[-1]

        for idx, row_xml in enumerate(data_rows):
            cells = row_xml.findall(qn('w:tc'))
            if idx < len(subject_notes):
                # Normalize row height to 315 dxa with hRule="atLeast"
                trPr = row_xml.find(qn('w:trPr'))
                if trPr is None:
                    trPr = row_xml.makeelement(qn('w:trPr'), {})
                    row_xml.insert(0, trPr)
                for h in trPr.findall(qn('w:trHeight')):
                    trPr.remove(h)
                trHeight = trPr.makeelement(qn('w:trHeight'), {
                    qn('w:val'): '315',
                    qn('w:hRule'): 'atLeast',
                })
                trPr.append(trHeight)

                subj_data = subject_notes[idx]
                cell_values = [
                    subj_data['code'],
                    subj_data['name'],
                    subj_data['note'],
                ]
                for ci, val in enumerate(cell_values):
                    if ci < len(cells):
                        for p in cells[ci].findall(qn('w:p')):
                            for r in p.findall(qn('w:r')):
                                t = r.find(qn('w:t'))
                                if t is not None:
                                    t.text = val
                                    break
                            else:
                                r_el = p.makeelement(qn('w:r'), {})
                                rpr_src = data_rows[0].findall(qn('w:tc'))[0] \
                                    .findall(qn('w:p'))[0].findall(qn('w:r'))
                                if rpr_src:
                                    rpr = rpr_src[0].find(qn('w:rPr'))
                                    if rpr is not None:
                                        r_el.append(deepcopy(rpr))
                                t_el = r_el.makeelement(qn('w:t'), {})
                                t_el.text = val
                                r_el.append(t_el)
                                p.append(r_el)
                            break
            else:
                tbl_xml.remove(row_xml)

        if len(subject_notes) > len(data_rows):
            ref_row = data_rows[0]
            for idx in range(len(data_rows), len(subject_notes)):
                subj_data = subject_notes[idx]
                new_row = deepcopy(ref_row)
                # Normalize row height to 315 dxa with hRule="atLeast"
                trPr = new_row.find(qn('w:trPr'))
                if trPr is None:
                    trPr = new_row.makeelement(qn('w:trPr'), {})
                    new_row.insert(0, trPr)
                for h in trPr.findall(qn('w:trHeight')):
                    trPr.remove(h)
                trHeight = trPr.makeelement(qn('w:trHeight'), {
                    qn('w:val'): '315',
                    qn('w:hRule'): 'atLeast',
                })
                trPr.append(trHeight)
                cells = new_row.findall(qn('w:tc'))
                cell_values = [
                    subj_data['code'],
                    subj_data['name'],
                    subj_data['note'],
                ]
                for ci, val in enumerate(cell_values):
                    if ci < len(cells):
                        for p in cells[ci].findall(qn('w:p')):
                            for r in p.findall(qn('w:r')):
                                t = r.find(qn('w:t'))
                                if t is not None:
                                    t.text = val
                                    break
                            break
                footer_row.addprevious(new_row)

        # Nota Media
        footer_cells = footer_row.findall(qn('w:tc'))
        nota_written = False
        for cell in reversed(footer_cells):
            cell_text = ''.join(
                t.text or ''
                for p in cell.findall(qn('w:p'))
                for r in p.findall(qn('w:r'))
                for t in r.findall(qn('w:t'))
            ).strip()
            if 'Nota Media' in cell_text or 'nota media' in cell_text.lower():
                continue
            for p in cell.findall(qn('w:p')):
                runs = p.findall(qn('w:r'))
                if runs:
                    for r in runs:
                        t = r.find(qn('w:t'))
                        if t is not None:
                            t.text = nota_media
                            nota_written = True
                            break
                else:
                    r_el = p.makeelement(qn('w:r'), {})
                    t_el = r_el.makeelement(qn('w:t'), {})
                    t_el.text = nota_media
                    r_el.append(t_el)
                    p.append(r_el)
                    nota_written = True
                break
            if nota_written:
                break

        if not nota_written and footer_cells:
            last_cell = footer_cells[-1]
            paras = last_cell.findall(qn('w:p'))
            target_p = paras[-1] if paras else None
            if target_p is not None:
                r_el = target_p.makeelement(qn('w:r'), {})
                t_el = r_el.makeelement(qn('w:t'), {})
                t_el.text = '  ' + nota_media
                r_el.append(t_el)
                target_p.append(r_el)

        if top_font_size:
            table_font_size = Pt(7.5) if is_physical else top_font_size
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for r in para.runs:
                                if r.font:
                                    r.font.size = table_font_size

        tmp_docx = tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False, prefix='cert_partial_'
        )
        doc.save(tmp_docx.name)
        tmp_docx.close()
        self._restore_vertical_legal_text(tpl_path, tmp_docx.name)
        if self.signer == 'raimon' and not is_physical:
            self._ensure_signature_logo(tmp_docx.name)
        if self.certificate_type in ('physical', 'physical_apostilled'):
            self._remove_header_logo(tmp_docx.name)
        else:
            self._ensure_bottom_right_arcs(tmp_docx.name)
        return tmp_docx.name
