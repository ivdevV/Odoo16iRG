# -*- coding: utf-8 -*-
import os
import hashlib
from zipfile import ZipFile
from lxml import etree

from docx import Document as DocxDocument

from odoo import fields
from odoo.tests.common import TransactionCase


class TestIrgCertificatePartial(TransactionCase):

    def _assert_vertical_legal_text_is_visible_in_xml(self, res_file):
        with ZipFile(res_file) as docx_zip:
            document_xml = ''.join(
                docx_zip.read(name).decode('utf-8', errors='ignore')
                for name in docx_zip.namelist()
                if name.endswith('.xml')
            )

        self.assertIn('B56488687', document_xml)
        self.assertIn('B-603323', document_xml)
        self.assertIn('w:val="10"', document_xml)

    def _assert_bottom_right_arcs_are_visible_in_xml(self, res_file):
        with ZipFile(res_file) as docx_zip:
            document_xml = docx_zip.read('word/document.xml').decode(
                'utf-8', errors='ignore'
            )
            rels_xml = docx_zip.read('word/_rels/document.xml.rels').decode(
                'utf-8', errors='ignore'
            )
            package_names = set(docx_zip.namelist())

        self.assertIn('name="Bottom Right Arcs"', document_xml)
        self.assertIn('behindDoc="1"', document_xml)
        self.assertIn('relativeFrom="page"><wp:align>right</wp:align>', document_xml)
        self.assertIn('relativeFrom="page"><wp:align>bottom</wp:align>', document_xml)
        self.assertIn('Target="media/bottom_right_arcs.png"', rels_xml)
        self.assertIn('word/media/bottom_right_arcs.png', package_names)

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        # Create student and course
        cls.partner = cls.env['res.partner'].create({'name': 'Test Student Partial'})
        cls.course = cls.env['op.course'].create({'name': 'Test Course Partial', 'code': 'TCPART01'})
        cls.batch = cls.env['op.batch'].create({
            'name': 'Batch Partial',
            'code': 'BPPART',
            'course_id': cls.course.id,
            'start_date': fields.Date.today(),
            'end_date': fields.Date.today(),
        })

        # Admissions
        cls.product = cls.env['product.product'].create({
            'name': 'Test Product Partial',
            'type': 'service',
        })
        cls.register = cls.env['op.admission.register'].create({
            'name': 'Test Register Partial',
            'course_id': cls.course.id,
            'start_date': fields.Date.today(),
            'end_date': fields.Date.today(),
            'min_count': 1,
            'max_count': 100,
            'product_id': cls.product.id,
        })
        cls.admission = cls.env['op.admission'].create({
            'name': 'ADM-TEST-PARTIAL',
            'partner_id': cls.partner.id,
            'course_id': cls.course.id,
            'register_id': cls.register.id,
            'gender': 'm',
            'first_name': 'Test',
            'last_name': 'Student',
            'email': 'test.partial@example.com',
            'birth_date': '1990-01-01',
        })

        # Gradebook template config
        cls.gradebook_tmpl = cls.env['app.gradebook'].create({
            'name': 'Gradebook Template Test',
            'gradebook_template_ids': [(0, 0, {
                'type': 'exam',
                'qty': 2,
                'weight': 100,
            })]
        })
        # Link template to course
        cls.course.write({'gradebook_id': cls.gradebook_tmpl.id})
        cls.tmpl_line = cls.gradebook_tmpl.gradebook_template_ids[0]

        # Gradebooks
        cls.gradebook = cls.env['app.gradebook.student'].create({
            'partner_id': cls.partner.id,
            'course_id': cls.course.id,
            'batch_id': cls.batch.id,
            'admission_id': cls.admission.id,
        })

        # Subjects
        cls.subject_normal = cls.env['op.subject'].create({
            'name': 'Subject Compulsory A',
            'code': 'SCA',
            'course_id': cls.course.id,
            'subject_type': 'compulsory',
        })
        cls.subject_pending = cls.env['op.subject'].create({
            'name': 'Subject Compulsory B',
            'code': 'SCB',
            'course_id': cls.course.id,
            'subject_type': 'compulsory',
        })

        # Link subjects to student gradebook
        cls.gb_subj_a = cls.env['app.gradebook.subject'].create({
            'gradebook_student_id': cls.gradebook.id,
            'op_subject_id': cls.subject_normal.id,
        })
        cls.gb_subj_b = cls.env['app.gradebook.subject'].create({
            'gradebook_student_id': cls.gradebook.id,
            'op_subject_id': cls.subject_pending.id,
        })

        # Create results for Subject A (2 exams, so it's complete)
        cls.env['app.gradebook.result'].create({
            'gradebook_subject_id': cls.gb_subj_a.id,
            'survey_type': 'exam',
            'scoring_total': 8.5,
        })
        cls.env['app.gradebook.result'].create({
            'gradebook_subject_id': cls.gb_subj_a.id,
            'survey_type': 'exam',
            'scoring_total': 9.5,
        })
        # Force computation
        cls.gb_subj_a.compute_final_subject_note()
        cls.gb_subj_a.compute_point_average()

        # Create only 1 result for Subject B (2 exams required, so it's incomplete / pending)
        cls.env['app.gradebook.result'].create({
            'gradebook_subject_id': cls.gb_subj_b.id,
            'survey_type': 'exam',
            'scoring_total': 7.0,
        })
        # Force computation
        cls.gb_subj_b.compute_final_subject_note()
        cls.gb_subj_b.compute_point_average()

    def test_01_partial_gradebook_fill_template(self):
        """Check template filling logic for partial gradebooks."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'document_type': 'gradebook_partial',
            'certificate_type': 'digital',
            'state': 'draft',
        })
        # Run template filling logic
        res_file = cert._fill_template()
        self.assertTrue(res_file)

    def test_02_partial_gradebook_first_sentence_has_bold_student_and_course(self):
        """Student and course names are bold in the first descriptive line."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'document_type': 'gradebook_partial',
            'certificate_type': 'digital',
            'state': 'draft',
        })

        res_file = cert._fill_template()
        self._assert_vertical_legal_text_is_visible_in_xml(res_file)
        self._assert_bottom_right_arcs_are_visible_in_xml(res_file)
        document = DocxDocument(res_file)
        paragraph = next(
            (
                para for para in document.paragraphs
                if 'Test Student Partial' in para.text
                and 'Test Course Partial' in para.text
            ),
            None,
        )

        self.assertIsNotNone(paragraph)
        expected_id_label = 'DNI/Pasaporte'
        if 'l10n_latam_identification_type_id' in self.partner._fields:
            identification_type = self.partner.l10n_latam_identification_type_id
            if identification_type:
                expected_id_label = identification_type.name
        self.assertIn(
            'Que Test Student Partial con %s  consta matriculado/a en '
            'el Test Course Partial durante el período académico' % expected_id_label,
            paragraph.text,
        )
        self.assertTrue(
            any(
                run.text == 'Test Student Partial' and run.bold is True
                for run in paragraph.runs
            )
        )
        self.assertTrue(
            any(
                run.text == 'Test Course Partial' and run.bold is True
                for run in paragraph.runs
            )
        )

    def test_03_partial_gradebook_dpto_intro_and_layout_are_adjusted(self):
        """Department signer uses requested issuer text and table-width layout."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'document_type': 'gradebook_partial',
            'certificate_type': 'digital',
            'signer': 'dpto_academico',
            'state': 'draft',
        })

        res_file = cert._fill_template()
        document = DocxDocument(res_file)

        intro = next(
            (
                para for para in document.paragraphs
                if 'El Instituto Raimon Gaja, con CIF B-56488687' in para.text
            ),
            None,
        )
        first_sentence = next(
            (
                para for para in document.paragraphs
                if 'Test Student Partial' in para.text
                and 'Test Course Partial' in para.text
            ),
            None,
        )
        closing = next(
            (
                para for para in document.paragraphs
                if 'Para que así conste' in para.text
            ),
            None,
        )
        signature_paragraph = next(
            (
                para for para in document.paragraphs
                if 'Departamento Académico' in para.text
                and 'Instituto Raimon Gaja' in para.text
            ),
            None,
        )

        self.assertIsNotNone(intro)
        self.assertEqual(
            intro.text,
            'El Instituto Raimon Gaja, con CIF B-56488687 en calle '
            'Córcega 213, 1º 2ª, 08036 Barcelona.',
        )
        self.assertIsNotNone(intro)
        self.assertEqual(intro.alignment, 0)
        self.assertEqual(intro.paragraph_format.left_indent.twips, -172)
        self.assertEqual(intro.paragraph_format.right_indent.twips, -783)
        for paragraph in (first_sentence, closing):
            self.assertIsNotNone(paragraph)
            self.assertEqual(paragraph.alignment, 3)
            self.assertEqual(paragraph.paragraph_format.left_indent.twips, -172)
            self.assertEqual(paragraph.paragraph_format.right_indent.twips, -783)
        self.assertIsNotNone(signature_paragraph)
        self.assertEqual(
            signature_paragraph.text.splitlines(),
            ['Departamento Académico', 'Instituto Raimon Gaja'],
        )
        self.assertEqual(signature_paragraph.alignment, 0)
        self.assertEqual(signature_paragraph.paragraph_format.left_indent.twips, -172)
        self.assertEqual(signature_paragraph.paragraph_format.right_indent.twips, -783)

        # Verificar que la firma (el dibujo) existe en el cuerpo en un párrafo vacío (inline)
        has_signature_drawing = False
        for p in document.paragraphs:
            drawings = p._element.xpath('.//w:drawing')
            if drawings and not p.text.strip():
                has_signature_drawing = True
                break
        self.assertTrue(
            has_signature_drawing,
            "La firma del departamento académico no se encuentra en un párrafo independiente vacío en el documento final."
        )

    def test_04_partial_gradebook_raimon_intro_certifica_and_signature_align_with_table(self):
        """Raimon signer header, CERTIFICA and signature use the same text grid."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'document_type': 'gradebook_partial',
            'certificate_type': 'digital',
            'signer': 'raimon',
            'state': 'draft',
        })

        res_file = cert._fill_template()
        document = DocxDocument(res_file)

        signer_intro = next(
            (
                para for para in document.paragraphs
                if 'Raimon Gaja Jaumeandreu' in para.text
                and 'Director General' in para.text
            ),
            None,
        )
        certifica = next(
            (para for para in document.paragraphs if para.text.strip() == 'CERTIFICA:'),
            None,
        )
        signature_name = next(
            (
                para for para in document.paragraphs
                if 'Raimon Gaja Jaumeandreu' in para.text
                and 'Instituto Raimon Gaja' in para.text
            ),
            None,
        )

        for paragraph in (signer_intro, certifica, signature_name):
            self.assertIsNotNone(paragraph)
            self.assertEqual(paragraph.paragraph_format.left_indent.twips, -172)
            self.assertEqual(paragraph.paragraph_format.right_indent.twips, -783)

    def test_05_partial_gradebook_all_pending_fill_template(self):
        """Check template filling logic when all compulsory subjects are pending."""
        # Unlink results for Subject A to make it pending too
        self.gb_subj_a.gradebook_result_ids.unlink()
        self.gb_subj_a.compute_final_subject_note()
        self.gb_subj_a.compute_point_average()

        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'document_type': 'gradebook_partial',
            'certificate_type': 'digital',
            'state': 'draft',
        })
        res_file = cert._fill_template()
        self.assertTrue(res_file)

    def test_06_partial_gradebook_uses_partner_identification_type_name(self):
        """Partial gradebook must print the exact partner identification type."""
        if 'l10n_latam_identification_type_id' not in self.partner._fields:
            self.skipTest('l10n_latam identification type field is not available.')

        identification_type = self.env.ref(
            'l10n_latam_base.it_vat', raise_if_not_found=False
        )
        if not identification_type:
            identification_type = self.env['l10n_latam.identification.type'].create({
                'name': 'VAT',
            })
        self.partner.with_context(no_vat_validation=True).write({
            'l10n_latam_identification_type_id': identification_type.id,
            'vat': 'ID-123456',
        })
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'document_type': 'gradebook_partial',
            'certificate_type': 'digital',
            'state': 'draft',
        })

        res_file = cert._fill_template()
        document = DocxDocument(res_file)
        paragraph_text = '\n'.join(para.text for para in document.paragraphs)

        self.assertIn('%s ID-123456' % identification_type.name, paragraph_text)

    def _assert_bottom_right_arcs_are_absent_in_xml(self, res_file):
        with ZipFile(res_file) as docx_zip:
            document_xml = docx_zip.read('word/document.xml').decode(
                'utf-8', errors='ignore'
            )
            rels_xml = docx_zip.read('word/_rels/document.xml.rels').decode(
                'utf-8', errors='ignore'
            )
            package_names = set(docx_zip.namelist())

        self.assertNotIn('name="Bottom Right Arcs"', document_xml)
        self.assertNotIn('Target="media/bottom_right_arcs.png"', rels_xml)
        self.assertNotIn('word/media/bottom_right_arcs.png', package_names)

    def _assert_header_logo_is_removed_in_xml(self, res_file):
        header_name = 'word/header1.xml'
        with ZipFile(res_file) as docx_zip:
            if header_name not in docx_zip.namelist():
                return
            header_xml = etree.fromstring(docx_zip.read(header_name))
        
        runs_with_drawings = []
        for r in header_xml.xpath('.//*[local-name()="r"]'):
            if r.xpath('.//*[local-name()="drawing" or local-name()="AlternateContent" or local-name()="pict"]'):
                runs_with_drawings.append(r)
        self.assertEqual(len(runs_with_drawings), 0, "Header logo drawings were not removed.")

    def _assert_header_logo_is_present_in_xml(self, res_file):
        header_name = 'word/header1.xml'
        with ZipFile(res_file) as docx_zip:
            if header_name not in docx_zip.namelist():
                self.fail("Header XML not found in document.")
            header_xml = etree.fromstring(docx_zip.read(header_name))
        
        runs_with_drawings = []
        for r in header_xml.xpath('.//*[local-name()="r"]'):
            if r.xpath('.//*[local-name()="drawing" or local-name()="AlternateContent" or local-name()="pict"]'):
                runs_with_drawings.append(r)
        self.assertTrue(len(runs_with_drawings) > 0, "Header logo drawings were removed but should be present.")

    def test_07_partial_physical_omits_logo_and_arcs(self):
        """Physical/apostilled partial gradebook certificates must omit the header logo and decorative arcs."""
        for cert_type in ('physical', 'physical_apostilled'):
            cert = self.env['irg.certificate.request'].create({
                'gradebook_student_id': self.gradebook.id,
                'document_type': 'gradebook_partial',
                'certificate_type': cert_type,
                'shipping_type': 'national',
                'state': 'draft',
            })
            res_file = cert._fill_template()
            self._assert_bottom_right_arcs_are_absent_in_xml(res_file)
            self._assert_header_logo_is_removed_in_xml(res_file)

    def test_08_partial_non_physical_retains_logo_and_arcs(self):
        """Digital/custom partial gradebook certificates must retain the header logo and decorative arcs."""
        for cert_type in ('digital', 'custom'):
            cert = self.env['irg.certificate.request'].create({
                'gradebook_student_id': self.gradebook.id,
                'document_type': 'gradebook_partial',
                'certificate_type': cert_type,
                'state': 'draft',
            })
            res_file = cert._fill_template()
            self._assert_bottom_right_arcs_are_visible_in_xml(res_file)
            self._assert_header_logo_is_present_in_xml(res_file)

    def _assert_table_font_sizes_match_top_font_size(self, res_file):
        doc = DocxDocument(res_file)
        top_font_size = None
        for para in doc.paragraphs:
            if para.text.strip():
                for r in para.runs:
                    if r.font and r.font.size:
                        top_font_size = r.font.size
                        break
            if top_font_size:
                break
        self.assertIsNotNone(top_font_size, "Could not find top font size in document.")
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            self.assertEqual(
                                run.font.size,
                                top_font_size,
                                "Table run font size does not match calculated top font size."
                            )

    def _assert_signature_logo_is_present_and_referenced(self, res_file):
        from odoo.modules.module import get_module_resource
        logo_path = get_module_resource('irg_gradebook_certificates', 'static', 'src', 'img', 'logodesgastado.png')
        self.assertTrue(logo_path and os.path.isfile(logo_path), "Expected signature logo does not exist in irg_gradebook_certificates module.")
        with open(logo_path, 'rb') as f:
            expected_md5 = hashlib.md5(f.read()).hexdigest()

        with ZipFile(res_file) as z:
            matching_media_file = None
            for name in z.namelist():
                if name.startswith('word/media/'):
                    data = z.read(name)
                    if hashlib.md5(data).hexdigest() == expected_md5:
                        matching_media_file = name
                        break
            self.assertIsNotNone(matching_media_file, "logodesgastado.png not found by MD5 hash inside ZIP media.")
            
            doc_xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
            rels_xml = z.read('word/_rels/document.xml.rels').decode('utf-8', errors='ignore')
            
        relative_target = matching_media_file.replace('word/', '')
        import re
        r_ids = re.findall(rf'Id="([^"]+)"[^>]+Target="{re.escape(relative_target)}"', rels_xml)
        self.assertTrue(len(r_ids) > 0, f"Relationship for {relative_target} not found in document.xml.rels")
        
        referenced = any(r_id in doc_xml for r_id in r_ids)
        self.assertTrue(referenced, f"Relationship ID(s) {r_ids} for signature logo not referenced in document.xml")

        # Validate that the signature logo is placed in the same paragraph containing the handwritten signature, having at least 2 drawing elements
        doc = DocxDocument(res_file)
        sig_rel_id = None
        for rel_id, rel in doc.part.rels.items():
            target = getattr(rel, 'target_ref', '').lower()
            if any(img in target for img in ('media/image2.jpg', 'media/image2.png', 'media/image2.jpeg')):
                sig_rel_id = rel_id
                break
        self.assertIsNotNone(sig_rel_id, "Handwritten signature rel_id not found")

        paragraphs_to_search = list(doc.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    paragraphs_to_search.extend(cell.paragraphs)

        sig_para = None
        for para in paragraphs_to_search:
            for r in para.runs:
                embed_nodes = r._r.xpath('.//*[@*[local-name()="embed" and .="%s"]]' % sig_rel_id)
                if embed_nodes:
                    sig_para = para
                    break
            if sig_para:
                break

        self.assertIsNotNone(sig_para, "Paragraph containing the handwritten signature not found")
        drawing_runs = [r for r in sig_para.runs if r._r.xpath('.//*[local-name()="drawing" or local-name()="pict"]')]
        self.assertTrue(len(drawing_runs) >= 2, "Paragraph containing the handwritten signature does not contain at least 2 drawings")

    def test_09_partial_table_font_sizes_match_top_font_size(self):
        """Verify that table cell font size matches the top font size for partial gradebooks."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'document_type': 'gradebook_partial',
            'certificate_type': 'digital',
            'state': 'draft',
        })
        res_file = cert._fill_template()
        self._assert_table_font_sizes_match_top_font_size(res_file)

        cert_phys = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'document_type': 'gradebook_partial',
            'certificate_type': 'physical',
            'shipping_type': 'national',
            'state': 'draft',
        })
        res_file_phys = cert_phys._fill_template()
        self._assert_table_font_sizes_match_top_font_size(res_file_phys)

    def test_10_partial_signature_logo_present_for_raimon_signer(self):
        """Verify that Raimon Gaja's signature logo is present and referenced in partial certificates."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'document_type': 'gradebook_partial',
            'certificate_type': 'digital',
            'signer': 'raimon',
            'state': 'draft',
        })
        res_file = cert._fill_template()
        self._assert_signature_logo_is_present_and_referenced(res_file)

    def _assert_table_data_row_heights_are_315_atleast(self, res_file):
        from docx.oxml.ns import qn
        doc = DocxDocument(res_file)
        table = doc.tables[0]
        rows = list(table.rows)
        self.assertTrue(len(rows) > 2, "Table should have header, data, and footer rows")
        
        data_rows = rows[1:-1]
        for idx, row in enumerate(data_rows):
            trPr = row._tr.find(qn('w:trPr'))
            self.assertIsNotNone(trPr, f"Row {idx+1} is missing trPr")
            trHeight = trPr.find(qn('w:trHeight'))
            self.assertIsNotNone(trHeight, f"Row {idx+1} is missing trHeight")
            
            val = trHeight.get(qn('w:val'))
            hRule = trHeight.get(qn('w:hRule'))
            
            self.assertEqual(val, '315', f"Row {idx+1} height val is {val}, expected '315'")
            self.assertEqual(hRule, 'atLeast', f"Row {idx+1} height hRule is {hRule}, expected 'atLeast'")

    def test_11_partial_table_data_row_heights_are_315_atleast(self):
        """Verify that all data rows in the grades table for partial gradebooks have height=315 dxa and hRule=atLeast."""
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'document_type': 'gradebook_partial',
            'certificate_type': 'digital',
            'state': 'draft',
        })
        res_file = cert._fill_template()
        self._assert_table_data_row_heights_are_315_atleast(res_file)

        cert_phys = self.env['irg.certificate.request'].create({
            'gradebook_student_id': self.gradebook.id,
            'document_type': 'gradebook_partial',
            'certificate_type': 'physical',
            'shipping_type': 'national',
            'state': 'draft',
        })
        res_file_phys = cert_phys._fill_template()
        self._assert_table_data_row_heights_are_315_atleast(res_file_phys)

    def test_12_partial_gradebook_course_id_4_ects_text(self):
        """Verify ECTS text and detailed text when course ID is 4 for partial gradebooks."""
        course_4 = self.env['op.course'].browse(4)
        if not course_4.exists():
            course_tmp = self.env['op.course'].create({
                'name': 'Special Course 4',
                'code': 'SC04',
            })
            self.env.cr.execute("UPDATE op_course SET id = 4 WHERE id = %s", (course_tmp.id,))
            self.env.registry.clear_cache()
            course_4 = self.env['op.course'].browse(4)
        
        course_4.write({'gradebook_id': self.gradebook_tmpl.id})
        
        batch_4 = self.env['op.batch'].search([('course_id', '=', course_4.id)], limit=1)
        if not batch_4:
            batch_4 = self.env['op.batch'].create({
                'name': 'Batch 4',
                'code': 'B4',
                'course_id': course_4.id,
                'start_date': fields.Date.today(),
                'end_date': fields.Date.today(),
            })
        
        register_4 = self.env['op.admission.register'].create({
            'name': 'Register 4',
            'course_id': course_4.id,
            'start_date': fields.Date.today(),
            'end_date': fields.Date.today(),
            'min_count': 1,
            'max_count': 100,
            'product_id': self.product.id,
        })
        
        admission_4 = self.env['op.admission'].create({
            'name': 'ADM-TEST-4',
            'partner_id': self.partner.id,
            'course_id': course_4.id,
            'register_id': register_4.id,
            'gender': 'm',
            'first_name': 'Test4',
            'last_name': 'Student4',
            'email': 'test4@example.com',
            'birth_date': '1990-01-01',
        })
        
        gradebook_4 = self.env['app.gradebook.student'].create({
            'partner_id': self.partner.id,
            'course_id': course_4.id,
            'batch_id': batch_4.id,
            'admission_id': admission_4.id,
        })
        
        self.env['app.gradebook.subject'].create({
            'gradebook_student_id': gradebook_4.id,
            'op_subject_id': self.subject_normal.id,
        })
        
        cert = self.env['irg.certificate.request'].create({
            'gradebook_student_id': gradebook_4.id,
            'certificate_type': 'digital',
            'document_type': 'gradebook_partial',
            'state': 'draft',
        })
        
        res_file = cert._fill_template()
        document = DocxDocument(res_file)
        
        second_body = next(
            (
                para for para in document.paragraphs
                if '120 ECTS' in para.text
            ),
            None,
        )
        self.assertIsNotNone(second_body, "120 ECTS text not found in generated document.")
        self.assertIn('120 ECTS, equivalentes a 3000 horas de estudio', second_body.text)


