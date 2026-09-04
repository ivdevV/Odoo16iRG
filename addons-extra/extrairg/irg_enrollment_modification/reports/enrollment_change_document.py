# -*- coding: utf-8 -*-
import logging
import os
import subprocess
import tempfile

from odoo import _, models
from odoo.exceptions import UserError
from odoo.modules.module import get_module_resource

_logger = logging.getLogger(__name__)

MODALITY_LABELS = {
    'Online': 'Online',
    'Presencial': 'Presencial',
    'Homeclass': 'Homeclass',
}


class EnrollmentChangeDocument(models.AbstractModel):
    _name = 'irg.enrollment.change.document'
    _description = 'Relleno de plantilla de modificación de matrícula'

    def _template_path(self):
        path = get_module_resource(
            'irg_enrollment_modification',
            'static/src/templates',
            'modificacion_matricula.docx',
        )
        if not path or not os.path.isfile(path):
            raise UserError(_('No se encontró la plantilla de modificación de matrícula.'))
        return path

    def build_docx_bytes(self, change, stage='request'):
        from docx import Document

        doc = Document(self._template_path())
        self._fill_header(doc, change)
        self._fill_change_rows(doc, change)
        self._fill_proposal(doc, change)
        if stage == 'final':
            self._fill_resolution(doc, change)
        else:
            self._clear_resolution(doc)
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False, prefix='enroll_chg_')
        tmp.close()
        doc.save(tmp.name)
        try:
            with open(tmp.name, 'rb') as handle:
                return handle.read()
        finally:
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

    def build_pdf_bytes(self, change):
        from docx import Document

        doc = Document(self._template_path())
        self._fill_header(doc, change)
        self._fill_change_rows(doc, change)
        self._fill_proposal(doc, change)
        self._fill_resolution(doc, change)
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False, prefix='enroll_chg_')
        tmp.close()
        doc.save(tmp.name)
        try:
            return self._convert_to_pdf(tmp.name)
        finally:
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

    def _fill_header(self, doc, change):
        student_name = change.student_id.name or ''
        date_txt = change.create_date.strftime('%d/%m/%Y') if change.create_date else ''
        team = ''
        if change.sale_order_id and change.sale_order_id.team_id:
            team = change.sale_order_id.team_id.name or ''
        if len(doc.tables) < 1:
            return
        header = doc.tables[0]
        self._replace_label(header.rows[0].cells[0], 'FECHA DE SOLICITUD:', date_txt)
        if len(header.rows[0].cells) > 1:
            self._replace_label(header.rows[0].cells[1], 'DELEGACIÓN:', team)
        if len(header.rows) > 1:
            self._replace_label(header.rows[1].cells[0], 'NOMBRE Y APELLIDOS:', student_name)

    def _fill_change_rows(self, doc, change):
        if len(doc.tables) < 2:
            return
        table = doc.tables[1]
        rows = table.rows
        if len(rows) > 1:
            self._fill_origin_dest_row(
                rows[1],
                change.change_course or change.change_batch,
                self._course_batch_label(change, 'origin'),
                self._course_batch_label(change, 'dest'),
            )
        if len(rows) > 2:
            self._fill_origin_dest_row(
                rows[2],
                change.change_modality,
                self._modality_label(change.origin_modality),
                self._modality_label(change.dest_modality),
            )
        if len(rows) > 3:
            self._fill_origin_dest_row(
                rows[3],
                change.change_year,
                change.origin_year_id.name or '',
                change.dest_year_id.name or '',
            )
        if len(rows) > 4 and change.change_payment:
            origin = change.origin_payment_mode_id.name or ''
            dest = change.dest_payment_mode_id.name or ''
            left = rows[4].cells[0]
            if len(left.paragraphs) > 1:
                left.paragraphs[1].text = 'Grupo de origen: %s' % origin
            right = rows[4].cells[1] if len(rows[4].cells) > 1 else left
            if right.paragraphs:
                right.paragraphs[0].text = 'Grupo de destino: %s' % dest

    def _fill_proposal(self, doc, change):
        author = change.create_uid.name or ''
        if len(doc.tables) < 3:
            return
        table = doc.tables[2]
        for row in table.rows:
            for cell in row.cells:
                self._replace_label(cell, 'PROPUESTA HECHA POR:', author)

    def _clear_resolution(self, doc):
        if len(doc.tables) < 3:
            return
        for cell in self._iter_cells(doc.tables[2]):
            for paragraph in cell.paragraphs:
                raw = paragraph.text.strip()
                if raw.startswith('X APROBADA'):
                    paragraph.text = 'APROBADA'

    def _fill_resolution(self, doc, change):
        resolver = (change.finance_user_id or change.academic_user_id).name or ''
        if len(doc.tables) < 3:
            return
        for cell in self._iter_cells(doc.tables[2]):
            self._replace_label(cell, 'RESOLUCIÓN HECHA POR:', resolver)
            for paragraph in cell.paragraphs:
                raw = paragraph.text.strip().upper()
                if 'APROBADA' in raw and 'DENEGADA' not in raw:
                    paragraph.text = 'X APROBADA'
                if 'ÁREA FINANCIERA' in raw or 'AREA FINANCIERA' in raw:
                    if change.finance_user_id:
                        paragraph.text = 'X ÁREA FINANCIERA'

    def _iter_cells(self, table):
        for row in table.rows:
            for cell in row.cells:
                yield cell

    def _fill_origin_dest_row(self, row, enabled, origin, dest):
        if not enabled:
            return
        for paragraph in row.cells[0].paragraphs:
            raw = paragraph.text.strip()
            if raw.startswith('Grupo de origen'):
                paragraph.text = 'Grupo de origen: %s' % origin
            elif raw.startswith('Grupo de destino'):
                paragraph.text = 'Grupo de destino: %s' % dest

    def _course_batch_label(self, change, side):
        enrollment = change.student_course_id
        current_course = enrollment.course_id.name if enrollment else ''
        current_batch = enrollment.batch_id.name if enrollment else ''
        if side == 'origin':
            course = change.origin_course_id.name if change.change_course else current_course
            batch = change.origin_batch_id.name if change.change_batch else current_batch
        else:
            course = change.dest_course_id.name if change.change_course else current_course
            batch = change.dest_batch_id.name if change.change_batch else current_batch
        return ' / '.join(part for part in (course, batch) if part)

    def _modality_label(self, value):
        return MODALITY_LABELS.get(value, value or '')

    def _replace_label(self, cell, label, value):
        for paragraph in cell.paragraphs:
            if label in paragraph.text:
                prefix = paragraph.text.split(label, 1)[0]
                paragraph.text = '%s%s %s' % (prefix, label, value or '')
                return True
        return False

    @staticmethod
    def _convert_to_pdf(docx_path):
        out_dir = os.path.dirname(docx_path)
        last_error = None
        for binary in ('libreoffice', 'soffice'):
            try:
                subprocess.run(
                    [
                        binary, '--headless', '--norestore',
                        '--convert-to', 'pdf',
                        '--outdir', out_dir,
                        docx_path,
                    ],
                    check=True,
                    timeout=60,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                )
                last_error = None
                break
            except FileNotFoundError:
                last_error = 'missing'
            except subprocess.CalledProcessError as exc:
                _logger.error('LibreOffice conversion failed: %s', exc.stderr)
                raise UserError(_('Error al convertir el documento a PDF. Revise el log.'))
        if last_error == 'missing':
            raise UserError(
                _('LibreOffice no está instalado en el servidor. '
                  'Ejecute: apt-get install -y libreoffice-writer')
            )

        pdf_path = docx_path.rsplit('.', 1)[0] + '.pdf'
        if not os.path.isfile(pdf_path):
            raise UserError(_('No se generó el archivo PDF.'))
        with open(pdf_path, 'rb') as handle:
            pdf_bytes = handle.read()
        try:
            os.unlink(pdf_path)
        except OSError:
            pass
        return pdf_bytes
