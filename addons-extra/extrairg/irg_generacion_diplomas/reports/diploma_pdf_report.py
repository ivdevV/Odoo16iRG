# -*- coding: utf-8 -*-
from odoo import models, api, modules, _
from odoo.exceptions import UserError
import docx
from docx import Document
import tempfile
import subprocess
from reportlab.lib.pagesizes import A4, A3, landscape
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader, simpleSplit
from babel.dates import format_date
import io
import os
import qrcode


class DiplomaReportPDF(models.AbstractModel):
    _name = 'report.irg_generacion_diplomas.diploma_pdf'
    _description = 'Diploma PDF Report'

    def _get_image_path(self, image_name):
        """Get full path to image in module's static folder"""
        return modules.get_module_resource('irg_generacion_diplomas', 'static/src/img', image_name)

    def _register_fonts(self):
        """Register Inter fonts if available, otherwise use Helvetica"""
        font_regular = 'Helvetica'
        font_bold = 'Helvetica-Bold'
        
        try:
            # Check Regular
            font_path = modules.get_module_resource('irg_generacion_diplomas', 'static/src/fonts', 'Inter-Regular.ttf')
            if not (font_path and os.path.exists(font_path)):
                 font_path = modules.get_module_resource('irg_generacion_diplomas', 'static/src/fonts', 'Inter-regular.ttf')
            
            if font_path and os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('Inter', font_path))
                font_regular = 'Inter'

            # Check Bold
            font_bold_path = modules.get_module_resource('irg_generacion_diplomas', 'static/src/fonts', 'Inter-Bold.ttf')
            if not (font_bold_path and os.path.exists(font_bold_path)):
                 font_bold_path = modules.get_module_resource('irg_generacion_diplomas', 'static/src/fonts', 'Inter-bold.ttf')
            
            if font_bold_path and os.path.exists(font_bold_path):
                pdfmetrics.registerFont(TTFont('Inter-Bold', font_bold_path))
                font_bold = 'Inter-Bold'
                
        except Exception as e:
            # Log error but fallback to Helvetica
            pass
            
        return font_regular, font_bold

    def _generate_qr(self, url, size=90):
        """Generate QR code image"""
        qr = qrcode.QRCode(version=1, box_size=10, border=1)
        qr.add_data(url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        
        # Convert to bytes
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)
        return ImageReader(buffer)

    def _draw_centered_text(self, c, text, y, font_name, font_size, page_width):
        """Draw text centered on page"""
        c.setFont(font_name, font_size)
        text_width = c.stringWidth(text, font_name, font_size)
        x = (page_width - text_width) / 2
        c.drawString(x, y, text)

    def _draw_text_in_column(self, c, text, x, y, width, font_name, font_size, align='center'):
        """Draw text within a column area"""
        c.setFont(font_name, font_size)
        text_width = c.stringWidth(text, font_name, font_size)
        
        if align == 'center':
            draw_x = x + (width - text_width) / 2
        elif align == 'left':
            draw_x = x
        else:
            draw_x = x + width - text_width
            
        c.drawString(draw_x, y, text)

    def _draw_wrapped_text_in_column(self, c, text, x, y, width, font_name, font_size, align='center', leading=None):
        """Draw text wrapping to new lines if needed. Returns new Y position."""
        if leading is None:
            leading = font_size * 1.2
            
        c.setFont(font_name, font_size)
        lines = simpleSplit(text, font_name, font_size, width)
        
        current_y = y
        for line in lines:
            self._draw_text_in_column(c, line, x, current_y, width, font_name, font_size, align)
            current_y -= leading
            
        return current_y

    def _fit_single_line_font_size(self, c, text, font_name, max_font_size, min_font_size, max_width, step=0.5):
        """Return the biggest font size that keeps text on a single line inside max_width."""
        if not text:
            return max_font_size

        font_size = max_font_size
        while font_size > min_font_size and c.stringWidth(text, font_name, font_size) > max_width:
            font_size -= step

        return max(font_size, min_font_size)

    def _normalize_catalan_course_name(self, course_name):
        """Normalize common accent differences for Catalan rendering.

        Also convert the Spanish conjunction " y " to Catalan " i " when it appears
        in the course title, since source data may be entered in Spanish but the
        Catalan version should use the correct word.
        """
        if not course_name:
            return course_name

        normalized = course_name
        normalized = normalized.replace("Máster", "Màster")
        normalized = normalized.replace("máster", "màster")
        normalized = normalized.replace("Master", "Màster")
        normalized = normalized.replace("master", "màster")
        normalized = normalized.replace("Salud", "Salut")
        normalized = normalized.replace("salud", "salut")
        # convert Spanish conjunction y to Catalan i when surrounded by spaces
        normalized = normalized.replace(" y ", " i ")
        normalized = normalized.replace(" Y ", " I ")
        return normalized

    @api.model
    def _get_report_values(self, docids, data=None):
        return {'data': data}

    def _replace_in_paragraph(self, paragraph, old, new):
        """Replace *old* with *new* across runs that Word may have split."""
        import re
        full = ''.join(r.text for r in paragraph.runs)
        if old not in full:
            return False
        
        if old == 'NombreAlumno>>':
            # Support flexible match for NombreAlumno>> (e.g. \ue097\ue097NombreAlumno>>)
            pattern = r'[^\w\s]*' + re.escape(old)
            full = re.sub(pattern, new, full)
        else:
            full = full.replace(old, new)
            
        if paragraph.runs:
            paragraph.runs[0].text = full
            for r in paragraph.runs[1:]:
                r.text = ''
        return True

    @api.model
    def generate_diploma_pdf(self, data, diploma_type='digital'):
        """Generate the diploma PDF using python-docx and LibreOffice"""
        filename = (
            'Plantilla Diplomas iRG Digital final.docx'
            if diploma_type == 'digital'
            else 'Plantilla Diploma fisico.docx'
        )
        template_path = modules.get_module_resource('irg_generacion_diplomas', 'static', filename)
        if not template_path or not os.path.exists(template_path):
            raise UserError(
                _('No se encuentra la plantilla Word en %s') % (template_path or filename)
            )

        temp_qr_path = None
        temp_docx_path = None
        pdf_path = None

        try:
            # Generate QR code
            qr_url = data.get('qr_url') or 'https://institutoraimongaja.com'
            qr = qrcode.QRCode(version=1, box_size=10, border=1)
            qr.add_data(qr_url)
            qr.make(fit=True)
            qr_img = qr.make_image(fill_color="black", back_color="white")

            temp_qr = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            temp_qr_path = temp_qr.name
            qr_img.save(temp_qr)
            temp_qr.close()

            doc = Document(template_path)

            if diploma_type == 'digital':
                replacements = {
                    '<<Mastercat>>': data.get('course_name_cat') or '',
                    '<<Master>>': data.get('course_name_es') or '',
                    'NombreAlumno>>': data.get('student_name') or '',
                    '<<fechacat>>': data.get('date_cat') or '',
                    '<<fecha>>': data.get('date_es') or '',
                    '<<registro>>': data.get('registry_number') or '',
                }
            else:
                replacements = {
                    '<<NombreCursoCat>>': data.get('course_name_cat') or '',
                    '<<NombreCurso>>': data.get('course_name_es') or '',
                    '<<NombreAlumno>>': data.get('student_name') or '',
                    '<<FechaExpedidoCat>>': data.get('date_cat') or '',
                    '<<FechaExpedido>>': data.get('date_es') or '',
                    'IRG-2026-0126': data.get('registry_number') or '',
                }

            # Helper to apply replacements to paragraphs
            def replace_in_paragraphs(paragraphs):
                for p in paragraphs:
                    if diploma_type == 'digital' and '<<Imagen_QR>>' in p.text:
                        p.text = ''
                        p.add_run().add_picture(temp_qr_path, width=docx.shared.Inches(1.2))
                        continue
                    
                    for old, new in replacements.items():
                        self._replace_in_paragraph(p, old, new)

            # Replace in main body paragraphs
            replace_in_paragraphs(doc.paragraphs)

            # Replace in headers/footers of all sections
            for section in doc.sections:
                replace_in_paragraphs(section.header.paragraphs)
                replace_in_paragraphs(section.footer.paragraphs)

            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_paragraphs(cell.paragraphs)

            # For physical: locate inline drawing and replace its image part blob
            if diploma_type == 'physical':
                with open(temp_qr_path, 'rb') as f:
                    qr_bytes = f.read()
                    
                found_qr = False
                for p in doc.paragraphs:
                    for run in p.runs:
                        blips = run._r.xpath('.//*[local-name()="blip"]')
                        if blips:
                            for blip in blips:
                                embed_id = None
                                for attr_name, attr_val in blip.items():
                                    if attr_name.endswith('}embed'):
                                        embed_id = attr_val
                                        break
                                if embed_id and embed_id in doc.part.related_parts:
                                    doc.part.related_parts[embed_id]._blob = qr_bytes
                                    found_qr = True

            # Save modified document to a temporary file
            temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False, prefix='diploma_')
            temp_docx_path = temp_docx.name
            doc.save(temp_docx_path)
            temp_docx.close()

            pdf_filename = os.path.basename(temp_docx_path).rsplit('.', 1)[0] + '.pdf'
            pdf_path = os.path.join('/tmp', pdf_filename)

            # Convert to PDF using LibreOffice
            try:
                subprocess.run(
                    [
                        'libreoffice', '--headless', '--norestore',
                        '--convert-to', 'pdf',
                        '--outdir', '/tmp',
                        temp_docx_path
                    ],
                    check=True,
                    timeout=60,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                )
            except Exception as exc:
                raise UserError(
                    _('Error al convertir el diploma a PDF usando LibreOffice: %s') % str(exc)
                )

            if not os.path.isfile(pdf_path):
                raise UserError(_('No se generó el archivo PDF del diploma.'))

            with open(pdf_path, 'rb') as f:
                pdf_bytes = f.read()

            return pdf_bytes

        finally:
            # Cleanup all temp files
            for path in (temp_docx_path, temp_qr_path, pdf_path):
                if path and os.path.exists(path):
                    try:
                        os.unlink(path)
                    except OSError:
                        pass
