# -*- coding: utf-8 -*-
import os
import tempfile
import logging
from docx import Document as DocxDocument

from odoo import models, fields, api, _
from odoo.exceptions import UserError, ValidationError

_logger = logging.getLogger(__name__)

class IrgCertificateRequest(models.Model):
    _inherit = 'irg.certificate.request'

    session_id = fields.Many2one(
        'op.session',
        string='Sesión de Clase en Directo',
        tracking=True,
    )

    def _validate_attendance_request(self, vals=None):
        """
        Validate that session_id is provided and the course/batch is HomeClass
        if the document type is 'attendance'.
        """
        for rec in self:
            doc_type = vals.get('document_type', rec.document_type) if vals else rec.document_type
            if doc_type == 'attendance':
                session_id = vals.get('session_id', rec.session_id.id) if vals else rec.session_id.id
                if not session_id:
                    raise ValidationError(_("La sesión es obligatoria para certificados de asistencia."))
                
                gradebook_student_id = vals.get('gradebook_student_id', rec.gradebook_student_id.id) if vals else rec.gradebook_student_id.id
                gradebook = self.env['app.gradebook.student'].sudo().browse(gradebook_student_id)
                if not gradebook:
                    raise ValidationError(_("Debe seleccionar una libreta académica válida."))
                
                has_homeclass_modality = any(m.code == 'homeclass' for m in gradebook.course_id.irg_modality_ids)
                has_hc_batch = False
                if gradebook.batch_id:
                    has_hc_batch = (gradebook.batch_id.code == 'HC') or ('HC' in (gradebook.batch_id.code or '').upper()) or ('HC' in (gradebook.batch_id.name or '').upper())
                
                if not (has_homeclass_modality or has_hc_batch):
                    raise ValidationError(_("El certificado de asistencia solo está disponible para cursos con modalidad HomeClass o grupos HC."))

    @api.model_create_multi
    def create(self, vals_list):
        records = super(IrgCertificateRequest, self).create(vals_list)
        for rec in records:
            rec._validate_attendance_request()
        return records

    def write(self, vals):
        res = super(IrgCertificateRequest, self).write(vals)
        self._validate_attendance_request(vals)
        return res

    def _get_template_path(self):
        if self.document_type == 'attendance':
            module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            tpl_dir = os.path.join(module_path, 'static', 'src', 'templates')
            signer_suffix = 'dpto' if self.signer == 'dpto_academico' else 'raimon'
            filename = f'Plantilla-certificado-asistencia-{signer_suffix}.docx'
            return os.path.join(tpl_dir, filename)
        return super()._get_template_path()

    def _fill_template(self):
        if self.document_type != 'attendance':
            return super()._fill_template()

        self.ensure_one()
        tpl_path = self._get_template_path()
        if not os.path.isfile(tpl_path):
            raise UserError(
                _('No se encuentra la plantilla Word en %s') % tpl_path
            )

        doc = DocxDocument(tpl_path)
        self._scale_document_fonts(doc, percent=75)

        # --- Collect data ---------------------------------------------------
        partner = self.partner_id
        id_label = (
            partner.l10n_latam_identification_type_id.name
            if partner.l10n_latam_identification_type_id
            else 'DNI/Pasaporte'
        )
        documento = '%s %s' % (id_label, partner.vat or '')

        # Fecha corta DD/MM/YYYY y fecha larga "25 de marzo de 2026"
        fecha = (
            self.request_date.strftime('%d/%m/%Y') if self.request_date else ''
        )
        if self.request_date:
            meses = {
                1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
                5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
                9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre',
            }
            dt = self.request_date
            fecha_larga = '%d de %s de %d' % (dt.day, meses[dt.month], dt.year)
        else:
            fecha_larga = ''

        # Session details
        session = self.session_id
        class_title = session.class_title or session.name or ''
        subject_name = session.subject_id.name or ''
        
        # Format session date in Spanish
        if session.start_datetime:
            dt_class = fields.Datetime.context_timestamp(self, session.start_datetime)
            meses = {
                1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
                5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
                9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre',
            }
            fecha_clase = '%d de %s de %d' % (dt_class.day, meses[dt_class.month], dt_class.year)
        else:
            fecha_clase = ''

        # Reemplazo específico para asistencia
        target_text = 'regularmente a las clases del programa académico <<nombreCurso>> en el periodo <<añoCurso>>'
        replacement_text = 'a la clase "%s" de la asignatura "%s" impartida el día %s' % (
            class_title,
            subject_name,
            fecha_clase
        )

        replacements = {
            '<<NombreAlumno>>': partner.name or '',
            '<<DocumentoIdentidad>>': documento,
            '<<fechaLarga>>': fecha_larga,
            '<<fecha>>': fecha,
            target_text: replacement_text,
            '<<nombreAlumno>>': partner.name or '',
            '<<documento>>': documento,
        }

        for para in doc.paragraphs:
            for old, new in replacements.items():
                self._replace_in_paragraph(para, old, new)
        for section in doc.sections:
            for para in section.header.paragraphs:
                for old, new in replacements.items():
                    self._replace_in_paragraph(para, old, new)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old, new in replacements.items():
                            self._replace_in_paragraph(para, old, new)

        # Save filled document to a temp file
        tmp_docx = tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False, prefix='cert_'
        )
        doc.save(tmp_docx.name)
        tmp_docx.close()
        return tmp_docx.name
