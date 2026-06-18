import sys
import os
import tempfile
from unittest.mock import MagicMock

# 1. Define dummy classes to replace Odoo classes
class DummyModel:
    pass

# Setup mock Odoo modules before importing anything else
odoo_exceptions = MagicMock()
odoo_exceptions.UserError = Exception
odoo_exceptions.ValidationError = Exception
sys.modules['odoo.exceptions'] = odoo_exceptions

odoo_module = MagicMock()
def mock_get_module_resource(module, *args):
    base_dir = "/Users/ivrogo/Workspace/Proyectos iRG/Odoo16iRG/addons-extra/extrairg"
    return os.path.join(base_dir, module, *args)
odoo_module.get_module_resource = mock_get_module_resource
sys.modules['odoo.modules'] = MagicMock()
sys.modules['odoo.modules.module'] = odoo_module

sys.modules['odoo'] = MagicMock()
sys.modules['odoo.http'] = MagicMock()
sys.modules['odoo.addons'] = MagicMock()
sys.modules['odoo.addons.irg_gradebook_certificates'] = MagicMock()
sys.modules['odoo.addons.irg_gradebook_certificates.controllers'] = MagicMock()
sys.modules['odoo.addons.irg_gradebook_certificates.controllers.portal'] = MagicMock()

# Inject odoo helpers
import odoo
odoo.models = MagicMock()
odoo.models.Model = DummyModel
odoo.fields = MagicMock()
odoo.api = MagicMock()
odoo._ = lambda x: x

from docx import Document
from docx.shared import Pt

# 2. Import model files directly
import importlib.util

def import_file(module_name, file_path):
    spec = importlib.util.spec_from_file_location(module_name, file_path)
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module

gradebook_model_path = "/Users/ivrogo/Workspace/Proyectos iRG/Odoo16iRG/addons-extra/extrairg/irg_gradebook_certificates/models/irg_certificate_request.py"
gradebook_mod = import_file("irg_gradebook_certificates.models.irg_certificate_request", gradebook_model_path)
GradebookRequest = gradebook_mod.IrgCertificateRequest

partial_model_path = "/Users/ivrogo/Workspace/Proyectos iRG/Odoo16iRG/addons-extra/extrairg/irg_certificate_partial/models/irg_certificate_request.py"
partial_mod = import_file("irg_certificate_partial.models.irg_certificate_request", partial_model_path)
PartialRequest = partial_mod.IrgCertificateRequest


class DummyRequest:
    pass

# Define standard mocks for request fields and relationships
def create_mock_request(cert_type, doc_type, signer='raimon'):
    req = DummyRequest()
    req.certificate_type = cert_type
    req.document_type = doc_type
    req.signer = signer
    
    # Partner mock
    partner = MagicMock()
    partner.name = "Rosmary Garzón González"
    partner.vat = "33750344"
    id_type = MagicMock()
    id_type.name = "Pasaporte"
    partner.l10n_latam_identification_type_id = id_type
    req.partner_id = partner
    
    # Course and Batch mock
    course = MagicMock()
    course.name = "Máster en Psicología Clínica Infantojuvenil"
    course.id = 12
    req.course_id = course
    
    batch = MagicMock()
    batch.start_date = MagicMock()
    batch.start_date.year = 2025
    
    gradebook_student = MagicMock()
    gradebook_student.total_final = 8.85
    gradebook_student.batch_id = batch
    req.gradebook_student_id = gradebook_student
    
    import datetime
    req.request_date = datetime.date(2026, 6, 17)
    
    # Mock methods
    req.ensure_one = lambda: None
    
    # Subjects mock
    subj1 = MagicMock()
    subj1.op_subject_id.code = "PI01"
    subj1.op_subject_id.name = "Apego y Estilos Parentales"
    subj1.final_subject_note = 9.20
    
    req._get_certificate_subjects = lambda: [subj1]
    
    # Constants from class
    req._GRADEBOOK_TEXT_INDENT = GradebookRequest._GRADEBOOK_TEXT_INDENT
    req._GRADEBOOK_TABLE_WIDTH = GradebookRequest._GRADEBOOK_TABLE_WIDTH
    req._GRADEBOOK_PAGE_TEXT_WIDTH = GradebookRequest._GRADEBOOK_PAGE_TEXT_WIDTH
    req._GRADEBOOK_TEXT_RIGHT_INDENT = GradebookRequest._GRADEBOOK_TEXT_RIGHT_INDENT
    req._DPTO_ACADEMICO_INTRO = GradebookRequest._DPTO_ACADEMICO_INTRO
    
    # Bind methods
    req._get_template_path = lambda: GradebookRequest._get_template_path(req)
    req._scale_document_fonts = GradebookRequest._scale_document_fonts
    req._replace_dpto_academico_intro = lambda doc: GradebookRequest._replace_dpto_academico_intro(req, doc)
    req._compact_gradebook_vertical_legal_text = GradebookRequest._compact_gradebook_vertical_legal_text
    req._replace_gradebook_description_paragraph = lambda doc, p, d, c, per, ects: GradebookRequest._replace_gradebook_description_paragraph(req, doc, p, d, c, per, ects)
    req._replace_in_paragraph = GradebookRequest._replace_in_paragraph
    req._format_gradebook_body_paragraph = lambda p, justify=True: GradebookRequest._format_gradebook_body_paragraph(req, p, justify)
    req._format_gradebook_static_paragraphs = lambda doc: GradebookRequest._format_gradebook_static_paragraphs(req, doc)
    req._format_gradebook_signature_paragraph = lambda p: GradebookRequest._format_gradebook_signature_paragraph(req, p)
    req._ensure_signature_logo = lambda docx: GradebookRequest._ensure_signature_logo(req, docx)
    req._remove_header_logo = lambda docx: GradebookRequest._remove_header_logo(req, docx)
    req._ensure_bottom_right_arcs = lambda docx: GradebookRequest._ensure_bottom_right_arcs(req, docx)
    req._restore_gradebook_vertical_legal_text = GradebookRequest._restore_gradebook_vertical_legal_text
    req._replace_paragraph_text_with_bold_segments = GradebookRequest._replace_paragraph_text_with_bold_segments
    
    req._fill_template = lambda: GradebookRequest._fill_template(req)
    
    return req

def create_mock_partial_request(cert_type, signer='raimon'):
    req = DummyRequest()
    req.certificate_type = cert_type
    req.document_type = 'gradebook_partial'
    req.signer = signer
    
    # Partner mock
    partner = MagicMock()
    partner.name = "Rosmary Garzón González"
    partner.vat = "33750344"
    id_type = MagicMock()
    id_type.name = "Pasaporte"
    partner.l10n_latam_identification_type_id = id_type
    req.partner_id = partner
    
    # Course and Batch mock
    course = MagicMock()
    course.name = "Máster en Psicología Clínica Infantojuvenil"
    course.id = 12
    req.course_id = course
    
    batch = MagicMock()
    batch.start_date = MagicMock()
    batch.start_date.year = 2025
    
    gradebook_student = MagicMock()
    gradebook_student.total_final = 8.85
    gradebook_student.batch_id = batch
    req.gradebook_student_id = gradebook_student
    
    import datetime
    req.request_date = datetime.date(2026, 6, 17)
    
    # Mock methods
    req.ensure_one = lambda: None
    
    # Subjects mock
    subj1 = MagicMock()
    subj1.op_subject_id.code = "PI01"
    subj1.op_subject_id.name = "Apego y Estilos Parentales"
    subj1.final_subject_note = 9.20
    
    subj1._get_gradebook_info = lambda s: {'exam': {'qty': 1}}
    exam_res = MagicMock()
    exam_res.survey_type = 'exam'
    
    class MockRecordset:
        def __init__(self, items):
            self.items = items
        def filtered(self, func):
            return MockRecordset([i for i in self.items if func(i)])
        def __len__(self):
            return len(self.items)
        def __bool__(self):
            return bool(self.items)
            
    subj1.gradebook_result_ids = MockRecordset([exam_res])
    req._get_certificate_subjects = lambda: [subj1]
    
    # Constants
    req._PARTIAL_TEXT_INDENT = PartialRequest._PARTIAL_TEXT_INDENT
    req._PARTIAL_TABLE_WIDTH = PartialRequest._PARTIAL_TABLE_WIDTH
    req._PARTIAL_PAGE_TEXT_WIDTH = PartialRequest._PARTIAL_PAGE_TEXT_WIDTH
    req._PARTIAL_TEXT_RIGHT_INDENT = PartialRequest._PARTIAL_TEXT_RIGHT_INDENT
    req._DPTO_ACADEMICO_INTRO = PartialRequest._DPTO_ACADEMICO_INTRO
    
    # Bind methods (inheriting from GradebookRequest where missing)
    req._get_template_path = lambda: PartialRequest._get_template_path(req)
    req._scale_document_fonts = GradebookRequest._scale_document_fonts
    req._replace_dpto_academico_intro = lambda doc: PartialRequest._replace_dpto_academico_intro(req, doc)
    req._compact_vertical_legal_text = PartialRequest._compact_vertical_legal_text
    req._replace_paragraph_text_with_bold_segments = PartialRequest._replace_paragraph_text_with_bold_segments
    req._replace_in_paragraph = GradebookRequest._replace_in_paragraph
    req._format_partial_body_paragraph = lambda p, justify=True: PartialRequest._format_partial_body_paragraph(req, p, justify)
    req._format_partial_closing_paragraphs = lambda doc: PartialRequest._format_partial_closing_paragraphs(req, doc)
    req._format_partial_static_paragraphs = lambda doc: PartialRequest._format_partial_static_paragraphs(req, doc)
    req._format_partial_signature_paragraph = lambda p: PartialRequest._format_partial_signature_paragraph(req, p)
    req._restore_vertical_legal_text = PartialRequest._restore_vertical_legal_text
    req._ensure_signature_logo = lambda docx: GradebookRequest._ensure_signature_logo(req, docx)
    req._remove_header_logo = lambda docx: GradebookRequest._remove_header_logo(req, docx)
    req._ensure_bottom_right_arcs = lambda docx: GradebookRequest._ensure_bottom_right_arcs(req, docx)
    req._fill_template = lambda: PartialRequest._fill_template(req)
    return req

def run_gradebook_validation():
    print("=== RUNNING GRADEBOOK CERTIFICATE VALIDATION ===")
    
    # Physical
    req_phys = create_mock_request('physical', 'gradebook')
    res_path = req_phys._fill_template()
    doc = Document(res_path)
    
    # Check 1: Margin top (72 Pt + 37.5 Pt = 109.5 Pt)
    print(f"Physical top margin (Pt): {doc.sections[0].top_margin.pt} (Expected: 109.5)")
    assert doc.sections[0].top_margin.pt == 109.5, "Physical top margin is incorrect"
    
    # Check 2: Outer text size
    body_runs = [r for p in doc.paragraphs if p.text.strip() for r in p.runs if r.font and r.font.size]
    for r in body_runs:
        print(f"Body run '{r.text[:20]}...' size: {r.font.size.pt} Pt (Expected: 9.0)")
        assert r.font.size.pt == 9.0, "Body text font size was scaled"
        
    # Check 3: Table text size
    table_runs = [r for t in doc.tables for row in t.rows for c in row.cells for p in c.paragraphs for r in p.runs if r.font and r.font.size]
    for r in table_runs:
        print(f"Table cell run '{r.text[:10]}...' size: {r.font.size.pt} Pt (Expected: 7.5)")
        assert r.font.size.pt == 7.5, "Table text size is not 7.5 Pt"
        
    # Check 4: Absence of signatures
    sig_rel_ids = []
    for rel_id, rel in doc.part.rels.items():
        target = getattr(rel, 'target_ref', '').lower()
        if any(img in target for img in ('media/image2.jpg', 'media/image2.png', 'media/image2.jpeg')):
            sig_rel_ids.append(rel_id)
            
    embeds = []
    if sig_rel_ids:
        for para in list(doc.paragraphs) + [p for t in doc.tables for row in t.rows for c in row.cells for p in c.paragraphs]:
            for r in para.runs:
                for rel_id in sig_rel_ids:
                    embeds.extend(r._r.xpath('.//*[@*[local-name()="embed" and .="%s"]]' % rel_id))
    print(f"Embedded signature shapes found: {len(embeds)} (Expected: 0)")
    assert len(embeds) == 0, "Signature/stamp images still embedded in document XML."
    
    # Check 5: Name replacement
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    print("Checking that 'Raimon Gaja Jaumeandreu' is replaced by 'Raimon Gaja'...")
    assert 'Raimon Gaja Jaumeandreu' not in full_text, "Raimon Gaja Jaumeandreu still found in text"
    assert 'Raimon Gaja' in full_text, "Raimon Gaja not found in text"
    
    # Check 6: Closing sentence and space after
    closing_paras = [p for p in doc.paragraphs if 'Para que así conste' in p.text]
    assert len(closing_paras) > 0, "Closing sentence not found"
    print(f"Closing sentence: '{closing_paras[0].text}'")
    assert closing_paras[0].text == "Para que así conste, firmo la presente en Barcelona, a fecha 17 de junio de 2026", "Closing sentence is incorrect"
    print(f"Closing space after (Pt): {closing_paras[0].paragraph_format.space_after.pt} (Expected: 48.0)")
    assert closing_paras[0].paragraph_format.space_after.pt == 48.0, "Closing space after is incorrect"
    
    # Check 7: Signature block text and font size
    sig_paras = [p for p in doc.paragraphs if 'Director General iRG' in p.text]
    assert len(sig_paras) > 0, "Signature block with Director General iRG not found"
    print(f"Signature block: '{sig_paras[0].text.replace(chr(10), ' | ')}'")
    assert sig_paras[0].text == "Raimon Gaja\nDirector General iRG", "Signature text is incorrect"
    
    os.unlink(res_path)
    print("GRADEBOOK VALIDATION PASSED SUCCESSFULLY!\n")
 
def run_partial_validation():
    print("=== RUNNING PARTIAL CERTIFICATE VALIDATION ===")
    
    # Physical
    req_phys = create_mock_partial_request('physical')
    res_path = req_phys._fill_template()
    doc = Document(res_path)
    
    # Check 1: Margin top (72 Pt + 37.5 Pt = 109.5 Pt)
    print(f"Physical partial top margin (Pt): {doc.sections[0].top_margin.pt} (Expected: 109.5)")
    assert doc.sections[0].top_margin.pt == 109.5, "Physical top margin is incorrect"
    
    # Check 2: Outer text size
    body_runs = [r for p in doc.paragraphs if p.text.strip() for r in p.runs if r.font and r.font.size]
    for r in body_runs:
        print(f"Body run '{r.text[:20]}...' size: {r.font.size.pt} Pt (Expected: 9.0)")
        assert r.font.size.pt == 9.0, "Body text font size was scaled"
        
    # Check 3: Table text size
    table_runs = [r for t in doc.tables for row in t.rows for c in row.cells for p in c.paragraphs for r in p.runs if r.font and r.font.size]
    for r in table_runs:
        print(f"Table cell run '{r.text[:10]}...' size: {r.font.size.pt} Pt (Expected: 7.5)")
        assert r.font.size.pt == 7.5, "Table text size is not 7.5 Pt"
        
    # Check 4: Absence of signatures
    sig_rel_ids = []
    for rel_id, rel in doc.part.rels.items():
        target = getattr(rel, 'target_ref', '').lower()
        if any(img in target for img in ('media/image2.jpg', 'media/image2.png', 'media/image2.jpeg')):
            sig_rel_ids.append(rel_id)
            
    embeds = []
    if sig_rel_ids:
        for para in list(doc.paragraphs) + [p for t in doc.tables for row in t.rows for c in row.cells for p in c.paragraphs]:
            for r in para.runs:
                for rel_id in sig_rel_ids:
                    embeds.extend(r._r.xpath('.//*[@*[local-name()="embed" and .="%s"]]' % rel_id))
    print(f"Embedded signature shapes found: {len(embeds)} (Expected: 0)")
    assert len(embeds) == 0, "Signature/stamp images still embedded in document XML."
    
    # Check 5: Name replacement
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    print("Checking that 'Raimon Gaja Jaumeandreu' is replaced by 'Raimon Gaja'...")
    assert 'Raimon Gaja Jaumeandreu' not in full_text, "Raimon Gaja Jaumeandreu still found in text"
    assert 'Raimon Gaja' in full_text, "Raimon Gaja not found in text"
    
    # Check 6: Closing sentence and space after
    closing_paras = [p for p in doc.paragraphs if 'Para que así conste' in p.text]
    assert len(closing_paras) > 0, "Closing sentence not found"
    print(f"Closing sentence: '{closing_paras[0].text}'")
    assert closing_paras[0].text == "Para que así conste, firmo la presente en Barcelona, a fecha 17 de junio de 2026", "Closing sentence is incorrect"
    print(f"Closing space after (Pt): {closing_paras[0].paragraph_format.space_after.pt} (Expected: 48.0)")
    assert closing_paras[0].paragraph_format.space_after.pt == 48.0, "Closing space after is incorrect"
    
    # Check 7: Signature block text and font size
    sig_paras = [p for p in doc.paragraphs if 'Director General iRG' in p.text]
    assert len(sig_paras) > 0, "Signature block with Director General iRG not found"
    print(f"Signature block: '{sig_paras[0].text.replace(chr(10), ' | ')}'")
    assert sig_paras[0].text == "Raimon Gaja\nDirector General iRG", "Signature text is incorrect"
    
    os.unlink(res_path)
    print("PARTIAL VALIDATION PASSED SUCCESSFULLY!\n")
if __name__ == '__main__':
    run_gradebook_validation()
    run_partial_validation()
